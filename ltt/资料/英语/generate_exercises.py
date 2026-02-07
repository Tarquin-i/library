#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
小学六年级英语练习题生成器
生成第2天（词汇基础）和第3天（句型基础）的练习题 Word 文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='宋体', font_size=12):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_title(doc, title, subtitle):
    """添加主标题和副标题"""
    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    set_chinese_font(run, '黑体', 18)

    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    set_chinese_font(run, '宋体', 14)
    doc.add_paragraph()  # 空行

def add_section_title(doc, title):
    """添加分节标题（带下划线）"""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.underline = True
    set_chinese_font(run, '黑体', 14)

def add_question_header(doc, number, title):
    """添加大题标题"""
    p = doc.add_paragraph()
    run = p.add_run(f'{number}、{title}')
    run.bold = True
    set_chinese_font(run, '宋体', 12)

def add_question(doc, text, indent=True):
    """添加题目"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_chinese_font(run, '宋体', 11)

def add_answer_section(doc, answers):
    """添加参考答案部分"""
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('参考答案')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    set_chinese_font(run, '黑体', 16)
    doc.add_paragraph()

    for section, ans_list in answers.items():
        p = doc.add_paragraph()
        run = p.add_run(section)
        run.bold = True
        set_chinese_font(run, '宋体', 12)
        for ans in ans_list:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(ans)
            set_chinese_font(run, '宋体', 11)


# ============ 第2天：词汇基础 题库 ============

DAY2_PHRASES_TRANSLATE = [
    ('寻找', 'look for'),
    ('看', 'look at'),
    ('照顾', 'look after'),
    ('起床', 'get up'),
    ('上床睡觉', 'go to bed'),
    ('醒来', 'wake up'),
    ('穿上', 'put on'),
    ('脱下', 'take off'),
    ('打开（电器）', 'turn on'),
    ('关闭（电器）', 'turn off'),
    ('在周末', 'on the weekend'),
    ('在晚上', 'at night'),
    ('在早上', 'in the morning'),
    ('在学校', 'at school'),
    ('准时', 'on time'),
    ('及时', 'in time'),
    ('回家', 'go home'),
    ('做作业', 'do homework'),
    ('吃早餐', 'have breakfast'),
    ('上学', 'go to school'),
]

DAY2_FILL_BLANKS = [
    ('I am _______ my keys. Have you seen them?', 'looking for', '我正在找我的钥匙。你看到了吗？'),
    ('Please _______ the baby while I cook dinner.', 'look after', '我做晚饭时请照顾宝宝。'),
    ('He usually _______ at 7 o\'clock every morning.', 'gets up', '他通常每天早上7点起床。'),
    ('It\'s cold outside. Please _______ your coat.', 'put on', '外面很冷。请穿上你的外套。'),
    ('Don\'t forget to _______ the lights before you leave.', 'turn off', '离开前别忘了关灯。'),
    ('She does her homework _______ every day.', 'at night', '她每天晚上做作业。'),
    ('We have PE class _______.', 'in the morning', '我们早上有体育课。'),
    ('Please _______ the blackboard.', 'look at', '请看黑板。'),
    ('I _______ at 6:30 and then have breakfast.', 'wake up', '我6:30醒来然后吃早餐。'),
    ('Can you _______ the TV? I want to watch the news.', 'turn on', '你能打开电视吗？我想看新闻。'),
    ('It\'s hot. Please _______ your jacket.', 'take off', '很热。请脱下你的夹克。'),
    ('We must arrive _______ for the meeting.', 'on time', '我们必须准时到达会议。'),
    ('The doctor came _______ to save the patient.', 'in time', '医生及时赶到救了病人。'),
    ('I _______ at 9 o\'clock every night.', 'go to bed', '我每晚9点上床睡觉。'),
    ('What do you usually do _______?', 'on the weekend', '你周末通常做什么？'),
]

DAY2_DISTINGUISH = [
    ('look at / look for / look after',
     [('Please _______ the picture on the wall.', 'look at'),
      ('I\'m _______ my pen. Where is it?', 'looking for'),
      ('Can you _______ my dog this weekend?', 'look after')]),
    ('put on / take off',
     [('It\'s raining. _______ your raincoat.', 'Put on'),
      ('It\'s warm inside. You can _______ your sweater.', 'take off')]),
    ('turn on / turn off',
     [('It\'s dark. Please _______ the light.', 'turn on'),
      ('Before you go to bed, _______ the computer.', 'turn off')]),
    ('on time / in time',
     [('The train arrived _______.', 'on time'),
      ('We got to the station just _______ to catch the train.', 'in time')]),
    ('get up / wake up',
     [('I _______ at 6:00 but I don\'t _______ until 6:30.', 'wake up, get up')]),
]

DAY2_COMPLETE = [
    ('我每天早上六点半起床。', 'I _______ _______ at 6:30 every morning.', 'get up'),
    ('请照顾好你的妹妹。', 'Please _______ _______ your little sister.', 'look after'),
    ('他正在找他的书包。', 'He is _______ _______ his schoolbag.', 'looking for'),
    ('上课前请关掉手机。', 'Please _______ _______ your phone before class.', 'turn off'),
    ('天冷了，穿上你的毛衣。', 'It\'s cold. _______ _______ your sweater.', 'Put on'),
    ('我们周末去公园。', 'We go to the park _______ _______ _______.', 'on the weekend'),
    ('她每天晚上九点上床睡觉。', 'She _______ _______ _______ at 9 p.m. every day.', 'goes to bed'),
    ('请看这张地图。', 'Please _______ _______ this map.', 'look at'),
    ('我早上七点醒来。', 'I _______ _______ at 7 a.m.', 'wake up'),
    ('进屋后请脱掉你的鞋子。', 'Please _______ _______ your shoes after entering the room.', 'take off'),
]


# ============ 第3天：句型基础 题库 ============

DAY3_SENTENCE_JUDGE = [
    ('I am a student.', '主系表', 'am 是 be 动词，a student 是表语'),
    ('She is happy.', '主系表', 'is 是 be 动词，happy 是表语'),
    ('They are teachers.', '主系表', 'are 是 be 动词，teachers 是表语'),
    ('He likes apples.', '主谓宾', 'likes 是实义动词，apples 是宾语'),
    ('We play football.', '主谓宾', 'play 是实义动词，football 是宾语'),
    ('She reads books every day.', '主谓宾', 'reads 是实义动词，books 是宾语'),
    ('The cat is on the table.', '主系表', 'is 是 be 动词，on the table 是表语'),
    ('I have a dog.', '主谓宾', 'have 是实义动词，a dog 是宾语'),
    ('Tom watches TV at night.', '主谓宾', 'watches 是实义动词，TV 是宾语'),
    ('My mother is a doctor.', '主系表', 'is 是 be 动词，a doctor 是表语'),
    ('The flowers are beautiful.', '主系表', 'are 是 be 动词，beautiful 是表语'),
    ('He does his homework.', '主谓宾', 'does 是实义动词，his homework 是宾语'),
    ('We are in the classroom.', '主系表', 'are 是 be 动词，in the classroom 是表语'),
    ('She eats breakfast at 7.', '主谓宾', 'eats 是实义动词，breakfast 是宾语'),
    ('The book is interesting.', '主系表', 'is 是 be 动词，interesting 是表语'),
]

DAY3_TO_QUESTION = [
    ('I am a student.', 'Are you a student?', 'be 动词提前'),
    ('She is happy.', 'Is she happy?', 'be 动词提前'),
    ('They are teachers.', 'Are they teachers?', 'be 动词提前'),
    ('He likes apples.', 'Does he like apples?', '加 Does，动词变原形'),
    ('We play football.', 'Do you play football?', '加 Do'),
    ('She reads books.', 'Does she read books?', '加 Does，动词变原形'),
    ('Tom watches TV.', 'Does Tom watch TV?', '加 Does，动词变原形'),
    ('I have a dog.', 'Do you have a dog?', '加 Do'),
    ('He is a doctor.', 'Is he a doctor?', 'be 动词提前'),
    ('They like music.', 'Do they like music?', '加 Do'),
    ('She goes to school.', 'Does she go to school?', '加 Does，动词变原形'),
    ('We are happy.', 'Are you happy?', 'be 动词提前'),
    ('He does homework.', 'Does he do homework?', '加 Does，动词变原形'),
    ('I am tired.', 'Are you tired?', 'be 动词提前'),
    ('She has a cat.', 'Does she have a cat?', '加 Does，动词变原形'),
]

DAY3_TO_NEGATIVE = [
    ('I am a student.', 'I am not a student.', 'be 动词后加 not'),
    ('She is happy.', 'She is not happy.', 'be 动词后加 not'),
    ('They are teachers.', 'They are not teachers.', 'be 动词后加 not'),
    ('He likes apples.', 'He does not like apples.', '加 does not，动词变原形'),
    ('We play football.', 'We do not play football.', '加 do not'),
    ('She reads books.', 'She does not read books.', '加 does not，动词变原形'),
    ('Tom watches TV.', 'Tom does not watch TV.', '加 does not，动词变原形'),
    ('I have a dog.', 'I do not have a dog.', '加 do not'),
    ('He is a doctor.', 'He is not a doctor.', 'be 动词后加 not'),
    ('They like music.', 'They do not like music.', '加 do not'),
    ('She goes to school.', 'She does not go to school.', '加 does not，动词变原形'),
    ('We are happy.', 'We are not happy.', 'be 动词后加 not'),
    ('He does homework.', 'He does not do homework.', '加 does not，动词变原形'),
    ('I am tired.', 'I am not tired.', 'be 动词后加 not'),
    ('She has a cat.', 'She does not have a cat.', '加 does not，动词变原形'),
]

DAY3_BE_DO_FILL = [
    ('_______ you a student? Yes, I _______.', 'Are, am', '主语 you 用 Are'),
    ('_______ she like apples? Yes, she _______.', 'Does, does', '第三人称单数用 Does'),
    ('He _______ not a teacher.', 'is', 'be 动词否定'),
    ('They _______ not play football.', 'do', '复数主语用 do'),
    ('_______ it cold today? Yes, it _______.', 'Is, is', '主语 it 用 Is'),
    ('She _______ not go to school on Sunday.', 'does', '第三人称单数用 does'),
    ('_______ your mother a doctor? Yes, she _______.', 'Is, is', '主语 your mother 用 Is'),
    ('_______ they have a car? No, they _______.', 'Do, don\'t', '复数主语用 Do'),
    ('I _______ not happy today.', 'am', '主语 I 用 am'),
    ('_______ Tom like music? Yes, he _______.', 'Does, does', '第三人称单数用 Does'),
    ('We _______ students.', 'are', '主语 We 用 are'),
    ('_______ you play basketball? Yes, I _______.', 'Do, do', '主语 you 用 Do'),
    ('The cat _______ on the table.', 'is', '主语 The cat 用 is'),
    ('_______ he do his homework? Yes, he _______.', 'Does, does', '第三人称单数用 Does'),
    ('My parents _______ not at home.', 'are', '复数主语用 are'),
]

DAY3_TRANSFORM = [
    ('He is a good boy.', '一般疑问句', 'Is he a good boy?'),
    ('She likes reading.', '否定句', 'She does not like reading.'),
    ('They are in the park.', '一般疑问句', 'Are they in the park?'),
    ('I play games every day.', '否定句', 'I do not play games every day.'),
    ('Tom has a new bike.', '一般疑问句', 'Does Tom have a new bike?'),
    ('We are happy.', '否定句', 'We are not happy.'),
    ('She watches TV at night.', '一般疑问句', 'Does she watch TV at night?'),
    ('He is my friend.', '否定句', 'He is not my friend.'),
    ('They do homework.', '一般疑问句', 'Do they do homework?'),
    ('I am a student.', '否定句', 'I am not a student.'),
    ('She is beautiful.', '一般疑问句', 'Is she beautiful?'),
    ('He plays football.', '否定句', 'He does not play football.'),
    ('We have lunch at 12.', '一般疑问句', 'Do you have lunch at 12?'),
    ('The dog is cute.', '否定句', 'The dog is not cute.'),
    ('They like swimming.', '一般疑问句', 'Do they like swimming?'),
]


# ============ 第4天：特殊疑问句+情景交际 题库 ============

# 疑问词选择填空（根据答句选择正确的疑问词）
DAY4_WH_FILL = [
    ('_______ is your name? My name is Tom.', 'What', '问"什么"用 what'),
    ('_______ are you from? I am from China.', 'Where', '问"哪里"用 where'),
    ('_______ is your birthday? It is on May 1st.', 'When', '问"什么时候"用 when'),
    ('_______ is that girl? She is my sister.', 'Who', '问"谁"用 who'),
    ('_______ are you late? Because I missed the bus.', 'Why', '问"为什么"用 why'),
    ('_______ do you go to school? I go to school by bus.', 'How', '问"怎样"用 how'),
    ('_______ is your favorite color? It is blue.', 'What', '问"什么"用 what'),
    ('_______ do you live? I live in Beijing.', 'Where', '问"哪里"用 where'),
    ('_______ do you get up? I get up at 6:30.', 'When', '问"什么时候"用 when'),
    ('_______ is your English teacher? Mr. Wang is.', 'Who', '问"谁"用 who'),
    ('_______ do you like apples? Because they are sweet.', 'Why', '问"为什么"用 why'),
    ('_______ old are you? I am twelve years old.', 'How', '问"多大"用 how old'),
    ('_______ is your phone number? It is 12345678.', 'What', '问"什么"用 what'),
    ('_______ is the library? It is next to the park.', 'Where', '问"哪里"用 where'),
    ('_______ does the movie start? It starts at 7 p.m.', 'When', '问"什么时候"用 when'),
]

# 疑问词意义匹配
DAY4_WH_MATCH = [
    ('What', '什么', '询问事物、名称、职业等'),
    ('When', '什么时候', '询问时间'),
    ('Where', '哪里', '询问地点、位置'),
    ('Who', '谁', '询问人物'),
    ('Why', '为什么', '询问原因'),
    ('How', '怎样/如何', '询问方式、状态、程度'),
    ('How old', '多大', '询问年龄'),
    ('How many', '多少（可数）', '询问可数名词数量'),
    ('How much', '多少（不可数）/多少钱', '询问不可数名词数量或价格'),
    ('How often', '多久一次', '询问频率'),
    ('How long', '多长/多久', '询问时间长度或物体长度'),
    ('How far', '多远', '询问距离'),
]

# 特殊疑问句语序练习（判断语序是否正确）
DAY4_WH_ORDER = [
    ('What is your name?', '正确', '疑问词 + be动词 + 主语'),
    ('Where you are from?', '错误，应为 Where are you from?', '疑问词 + be动词 + 主语'),
    ('When do you get up?', '正确', '疑问词 + do/does + 主语 + 动词原形'),
    ('Who is she?', '正确', '疑问词 + be动词 + 主语'),
    ('Why you are late?', '错误，应为 Why are you late?', '疑问词 + be动词 + 主语'),
    ('How does he go to school?', '正确', '疑问词 + do/does + 主语 + 动词原形'),
    ('What do you like?', '正确', '疑问词 + do/does + 主语 + 动词原形'),
    ('Where does she live?', '正确', '疑问词 + do/does + 主语 + 动词原形'),
    ('How old you are?', '错误，应为 How old are you?', '疑问词 + be动词 + 主语'),
    ('Who your teacher is?', '错误，应为 Who is your teacher?', '疑问词 + be动词 + 主语'),
    ('When is your birthday?', '正确', '疑问词 + be动词 + 主语'),
    ('Why does she cry?', '正确', '疑问词 + do/does + 主语 + 动词原形'),
    ('How many books do you have?', '正确', '疑问词 + 名词 + do/does + 主语 + 动词原形'),
    ('What time it is?', '错误，应为 What time is it?', '疑问词 + be动词 + 主语'),
    ('How much is this book?', '正确', '疑问词 + be动词 + 主语'),
]

# 情景交际 - 购物场景
DAY4_SHOPPING = [
    ('A: Can I help you?\nB: _______', 'Yes, I want to buy a book. / I\'m looking for a T-shirt.', '店员询问，顾客说明需求'),
    ('A: How much is this pen?\nB: _______', 'It\'s five yuan. / It\'s 5 dollars.', '询问价格，回答具体金额'),
    ('A: What color do you want?\nB: _______', 'I want a blue one. / Blue, please.', '询问颜色，回答颜色偏好'),
    ('A: What size do you need?\nB: _______', 'Size M, please. / I need a small one.', '询问尺码，回答尺码'),
    ('A: Here you are.\nB: _______', 'Thank you. / Thanks a lot.', '递东西时的礼貌回应'),
    ('A: _______\nB: It\'s 20 yuan.', 'How much is it? / What\'s the price?', '根据答句推断问句'),
    ('A: Do you have a smaller one?\nB: _______', 'Yes, here you are. / Sorry, this is the smallest.', '询问是否有其他尺码'),
    ('A: I\'ll take it.\nB: _______', 'OK. That\'s 30 yuan. / Great. Cash or card?', '决定购买后的回应'),
]

# 情景交际 - 问路场景
DAY4_ASKING_WAY = [
    ('A: Excuse me, where is the hospital?\nB: _______', 'Go straight and turn left. / It\'s next to the bank.', '询问地点，给出方向'),
    ('A: How can I get to the park?\nB: _______', 'You can take Bus No. 5. / Walk along this street.', '询问如何到达，给出方式'),
    ('A: Is the library far from here?\nB: _______', 'No, it\'s just 5 minutes\' walk. / Yes, you\'d better take a bus.', '询问距离远近'),
    ('A: _______\nB: Go straight and turn right at the corner.', 'Excuse me, how can I get to the post office?', '根据答句推断问句'),
    ('A: Excuse me, is there a bank near here?\nB: _______', 'Yes, there is one on Green Street. / Sorry, I don\'t know.', '询问附近是否有某地点'),
    ('A: Thank you very much.\nB: _______', 'You\'re welcome. / My pleasure.', '感谢后的礼貌回应'),
    ('A: Which bus should I take?\nB: _______', 'You can take Bus No. 10. / Take the subway Line 2.', '询问乘坐哪路车'),
    ('A: How far is it from here?\nB: _______', 'It\'s about 2 kilometers. / About 10 minutes by bus.', '询问具体距离'),
]

# 情景交际 - 问候场景
DAY4_GREETING = [
    ('A: Hello! How are you?\nB: _______', 'I\'m fine, thank you. And you? / Very well, thanks.', '日常问候'),
    ('A: Good morning!\nB: _______', 'Good morning! / Morning!', '早上问候'),
    ('A: Nice to meet you.\nB: _______', 'Nice to meet you, too.', '初次见面'),
    ('A: How do you do?\nB: _______', 'How do you do?', '正式问候，原句回答'),
    ('A: What\'s your name?\nB: _______', 'My name is... / I\'m...', '询问姓名'),
    ('A: Where are you from?\nB: _______', 'I\'m from China. / I come from Beijing.', '询问来自哪里'),
    ('A: See you tomorrow!\nB: _______', 'See you! / Bye-bye!', '告别'),
    ('A: Have a nice day!\nB: _______', 'You too! / Thank you, you too!', '祝福语回应'),
    ('A: How is your family?\nB: _______', 'They are fine, thank you. / Everyone is good.', '询问家人情况'),
    ('A: Long time no see!\nB: _______', 'Yes, it\'s been a while! / I missed you!', '久别重逢'),
]

# 综合练习 - 特殊疑问句造句
DAY4_MAKE_SENTENCE = [
    ('what / name / your / is', 'What is your name?', '用 what 询问姓名'),
    ('where / you / do / live', 'Where do you live?', '用 where 询问住址'),
    ('when / birthday / is / your', 'When is your birthday?', '用 when 询问生日'),
    ('who / teacher / your / is', 'Who is your teacher?', '用 who 询问人物'),
    ('why / late / you / are', 'Why are you late?', '用 why 询问原因'),
    ('how / to / school / do / go / you', 'How do you go to school?', '用 how 询问方式'),
    ('how old / you / are', 'How old are you?', '用 how old 询问年龄'),
    ('how many / books / have / you / do', 'How many books do you have?', '用 how many 询问数量'),
    ('how much / this / is / pen', 'How much is this pen?', '用 how much 询问价格'),
    ('what time / it / is', 'What time is it?', '用 what time 询问时间'),
    ('how often / exercise / do / you', 'How often do you exercise?', '用 how often 询问频率'),
    ('how far / school / is / your', 'How far is your school?', '用 how far 询问距离'),
]

# 综合练习 - 补全对话
DAY4_DIALOGUE_COMPLETE = [
    # 购物对话
    ('购物场景',
     '''A: Good morning! (1)_______
B: Yes, please. I want to buy a schoolbag.
A: (2)_______ do you like?
B: I like the blue one.
A: Here you are. It\'s 50 yuan.
B: (3)_______. Here is the money.
A: Thank you. (4)_______!''',
     ['Can I help you?', 'What color', 'I\'ll take it', 'Have a nice day']),

    # 问路对话
    ('问路场景',
     '''A: Excuse me. (1)_______?
B: Go straight and turn left at the traffic lights.
A: (2)_______?
B: No, it\'s not far. About 5 minutes\' walk.
A: (3)_______
B: You\'re welcome.''',
     ['Where is the bookstore / How can I get to the bookstore', 'Is it far from here', 'Thank you very much']),

    # 问候对话
    ('问候场景',
     '''A: Hello! (1)_______?
B: I\'m fine, thank you. And you?
A: I\'m fine, too. (2)_______?
B: My name is Li Ming.
A: Nice to meet you, Li Ming.
B: (3)_______''',
     ['How are you', 'What\'s your name', 'Nice to meet you, too']),

    # 综合对话1
    ('综合场景1',
     '''A: Good afternoon!
B: (1)_______!
A: (2)_______ is your English teacher?
B: Miss Wang is.
A: (3)_______ is she from?
B: She is from Shanghai.
A: (4)_______ does she teach English?
B: Because she loves English very much.''',
     ['Good afternoon', 'Who', 'Where', 'Why']),

    # 综合对话2
    ('综合场景2',
     '''A: Hi, Tom! (1)_______ do you usually get up?
B: I usually get up at 6:30.
A: (2)_______ do you go to school?
B: I go to school by bike.
A: (3)_______ is your school from your home?
B: It\'s about 2 kilometers.''',
     ['When / What time', 'How', 'How far']),
]


# ============ 第5天：语法专项（一般现在时）题库 ============

# 动词变形练习（原形→第三人称单数）
DAY5_VERB_CHANGE = [
    ('play', 'plays', '直接加 s'),
    ('read', 'reads', '直接加 s'),
    ('like', 'likes', '直接加 s'),
    ('eat', 'eats', '直接加 s'),
    ('drink', 'drinks', '直接加 s'),
    ('watch', 'watches', '以 ch 结尾加 es'),
    ('wash', 'washes', '以 sh 结尾加 es'),
    ('teach', 'teaches', '以 ch 结尾加 es'),
    ('go', 'goes', '以 o 结尾加 es'),
    ('do', 'does', '以 o 结尾加 es'),
    ('fix', 'fixes', '以 x 结尾加 es'),
    ('pass', 'passes', '以 s 结尾加 es'),
    ('study', 'studies', '辅音+y 结尾，变 y 为 i 加 es'),
    ('fly', 'flies', '辅音+y 结尾，变 y 为 i 加 es'),
    ('carry', 'carries', '辅音+y 结尾，变 y 为 i 加 es'),
    ('cry', 'cries', '辅音+y 结尾，变 y 为 i 加 es'),
    ('try', 'tries', '辅音+y 结尾，变 y 为 i 加 es'),
    ('have', 'has', '不规则变化'),
    ('play', 'plays', '元音+y 结尾，直接加 s'),
    ('say', 'says', '元音+y 结尾，直接加 s'),
]

# 时态判断题（判断是否为一般现在时）
DAY5_TENSE_JUDGE = [
    ('He plays football every day.', '是', '有 every day，表示经常性动作'),
    ('She is reading a book now.', '否', '有 now，是现在进行时'),
    ('They often go to school by bus.', '是', '有 often，表示习惯性动作'),
    ('I am watching TV.', '否', '有 am + doing，是现在进行时'),
    ('Tom usually gets up at 6:00.', '是', '有 usually，表示习惯性动作'),
    ('We are playing games.', '否', '有 are + doing，是现在进行时'),
    ('She always helps her mother.', '是', '有 always，表示经常性动作'),
    ('Look! The bird is flying.', '否', '有 Look! 和 is + doing，是现在进行时'),
    ('My father works in a hospital.', '是', '描述职业，表示客观事实'),
    ('The sun rises in the east.', '是', '描述自然规律，用一般现在时'),
    ('I like apples very much.', '是', '表示喜好，用一般现在时'),
    ('Listen! Someone is singing.', '否', '有 Listen! 和 is + doing，是现在进行时'),
    ('Water boils at 100°C.', '是', '描述科学事实，用一般现在时'),
    ('He is doing his homework now.', '否', '有 now 和 is + doing，是现在进行时'),
    ('She never eats meat.', '是', '有 never，表示习惯'),
]

# 动词形式选择题
DAY5_VERB_CHOOSE = [
    ('He _______ (play/plays) basketball every Sunday.', 'plays', '主语 He 是第三人称单数'),
    ('They _______ (go/goes) to school by bike.', 'go', '主语 They 是复数'),
    ('My mother _______ (cook/cooks) dinner every day.', 'cooks', '主语 My mother 是第三人称单数'),
    ('I _______ (like/likes) reading books.', 'like', '主语 I 用动词原形'),
    ('She _______ (have/has) a beautiful dress.', 'has', '主语 She 是第三人称单数，have→has'),
    ('Tom and Jerry _______ (watch/watches) TV together.', 'watch', '主语是两个人，用复数'),
    ('The dog _______ (run/runs) very fast.', 'runs', '主语 The dog 是第三人称单数'),
    ('We _______ (do/does) homework after school.', 'do', '主语 We 是复数'),
    ('My father _______ (drive/drives) to work.', 'drives', '主语 My father 是第三人称单数'),
    ('You _______ (speak/speaks) English very well.', 'speak', '主语 You 用动词原形'),
    ('The cat _______ (sleep/sleeps) on the sofa.', 'sleeps', '主语 The cat 是第三人称单数'),
    ('Birds _______ (fly/flies) in the sky.', 'fly', '主语 Birds 是复数'),
    ('She _______ (study/studies) hard every day.', 'studies', '主语 She 是第三人称单数，study→studies'),
    ('He _______ (do/does) his homework carefully.', 'does', '主语 He 是第三人称单数，do→does'),
    ('My parents _______ (work/works) in Beijing.', 'work', '主语 My parents 是复数'),
]

# 变形规则分类题
DAY5_RULE_CLASSIFY = [
    ('直接加 s', ['play', 'read', 'like', 'eat', 'drink', 'run', 'swim', 'write', 'sing', 'think']),
    ('以 s/x/ch/sh/o 结尾加 es', ['watch', 'wash', 'teach', 'go', 'do', 'fix', 'pass', 'brush', 'catch', 'miss']),
    ('辅音+y 结尾，变 y 为 i 加 es', ['study', 'fly', 'carry', 'cry', 'try', 'worry', 'hurry']),
    ('元音+y 结尾，直接加 s', ['play', 'say', 'stay', 'enjoy', 'buy']),
    ('不规则变化', ['have→has', 'be→is/am/are']),
]

# 时间标志词识别题
DAY5_TIME_WORDS = [
    ('always', '总是', '一般现在时标志词'),
    ('usually', '通常', '一般现在时标志词'),
    ('often', '经常', '一般现在时标志词'),
    ('sometimes', '有时', '一般现在时标志词'),
    ('never', '从不', '一般现在时标志词'),
    ('every day', '每天', '一般现在时标志词'),
    ('every week', '每周', '一般现在时标志词'),
    ('every month', '每月', '一般现在时标志词'),
    ('every year', '每年', '一般现在时标志词'),
    ('on Sundays', '在周日（每周日）', '一般现在时标志词'),
    ('in the morning', '在早上', '一般现在时标志词'),
    ('at night', '在晚上', '一般现在时标志词'),
]

# 不规则变化专练
DAY5_IRREGULAR = [
    ('have', 'has', 'She _______ a cat. (have)', 'has'),
    ('do', 'does', 'He _______ his homework every day. (do)', 'does'),
    ('go', 'goes', 'Tom _______ to school by bus. (go)', 'goes'),
    ('be', 'is/am/are', 'She _______ a student. (be)', 'is'),
    ('have', 'has', 'My brother _______ many books. (have)', 'has'),
    ('do', 'does', 'What _______ she do on weekends? (do)', 'does'),
    ('go', 'goes', 'The bus _______ to the park. (go)', 'goes'),
    ('be', 'is/am/are', 'They _______ happy. (be)', 'are'),
    ('have', 'has', 'The dog _______ a long tail. (have)', 'has'),
    ('do', 'does', '_______ he like apples? (do)', 'Does'),
]

# 时态填空题
DAY5_FILL_BLANK = [
    ('My sister _______ (watch) TV every evening.', 'watches', '第三人称单数 + every evening'),
    ('They _______ (play) football after school.', 'play', '复数主语用原形'),
    ('He _______ (go) to bed at 9 o\'clock.', 'goes', '第三人称单数，go→goes'),
    ('I _______ (get) up at 6:30 every morning.', 'get', '主语 I 用原形'),
    ('She _______ (study) English very hard.', 'studies', '第三人称单数，study→studies'),
    ('The earth _______ (move) around the sun.', 'moves', '科学事实用一般现在时'),
    ('Tom often _______ (help) his mother.', 'helps', '第三人称单数 + often'),
    ('We usually _______ (have) lunch at 12:00.', 'have', '复数主语用原形'),
    ('My father _______ (work) in a factory.', 'works', '第三人称单数'),
    ('_______ your brother _______ (like) music?', 'Does, like', '第三人称单数疑问句'),
    ('She _______ (not go) to school on Sundays.', 'doesn\'t go', '第三人称单数否定句'),
    ('The children _______ (be) very happy.', 'are', '复数主语用 are'),
    ('Water _______ (freeze) at 0°C.', 'freezes', '科学事实用一般现在时'),
    ('He never _______ (eat) breakfast.', 'eats', '第三人称单数 + never'),
    ('My parents always _______ (tell) me to study hard.', 'tell', '复数主语用原形'),
]

# 单句改错题
DAY5_CORRECT_ERROR = [
    ('He go to school every day.', 'go → goes', '主语 He 是第三人称单数'),
    ('She have a new bike.', 'have → has', '主语 She 是第三人称单数，have→has'),
    ('Tom watchs TV every night.', 'watchs → watches', 'watch 以 ch 结尾，加 es'),
    ('My mother cook dinner every day.', 'cook → cooks', '主语 My mother 是第三人称单数'),
    ('He studys English hard.', 'studys → studies', 'study 辅音+y 结尾，变 y 为 i 加 es'),
    ('The dog run very fast.', 'run → runs', '主语 The dog 是第三人称单数'),
    ('She gos to the park on Sundays.', 'gos → goes', 'go 以 o 结尾，加 es'),
    ('He dos his homework carefully.', 'dos → does', 'do 以 o 结尾，加 es'),
    ('My sister crys every night.', 'crys → cries', 'cry 辅音+y 结尾，变 y 为 i 加 es'),
    ('Tom and Mary plays football.', 'plays → play', '主语是两个人，用复数'),
    ('She don\'t like apples.', 'don\'t → doesn\'t', '第三人称单数否定用 doesn\'t'),
    ('Do he go to school by bus?', 'Do → Does', '第三人称单数疑问用 Does'),
    ('The sun rise in the east.', 'rise → rises', '主语 The sun 是第三人称单数'),
    ('He flys a kite in the park.', 'flys → flies', 'fly 辅音+y 结尾，变 y 为 i 加 es'),
    ('She washs her face every morning.', 'washs → washes', 'wash 以 sh 结尾，加 es'),
]

# 句子补全题
DAY5_COMPLETE_SENT = [
    ('他每天早上六点起床。', 'He _______ _______ at 6:00 every morning.', 'gets up'),
    ('她经常帮助她妈妈。', 'She often _______ her mother.', 'helps'),
    ('我爸爸在医院工作。', 'My father _______ in a hospital.', 'works'),
    ('太阳从东方升起。', 'The sun _______ in the east.', 'rises'),
    ('汤姆每天做作业。', 'Tom _______ his homework every day.', 'does'),
    ('她有一只可爱的猫。', 'She _______ a cute cat.', 'has'),
    ('他们通常乘公交车去学校。', 'They usually _______ to school by bus.', 'go'),
    ('玛丽每天晚上看电视。', 'Mary _______ TV every evening.', 'watches'),
    ('我弟弟非常努力学习。', 'My brother _______ very hard.', 'studies'),
    ('水在100度沸腾。', 'Water _______ at 100°C.', 'boils'),
    ('她从不吃早餐。', 'She never _______ breakfast.', 'eats'),
    ('他每周日踢足球。', 'He _______ football every Sunday.', 'plays'),
    ('我妈妈每天做饭。', 'My mother _______ dinner every day.', 'cooks'),
    ('鸟儿在天空飞翔。', 'Birds _______ in the sky.', 'fly'),
    ('他总是按时到校。', 'He always _______ to school on time.', 'gets'),
]


# ============ 第6天：语法专项（一般过去时+一般将来时）题库 ============

# 一般过去时标志词
DAY6_PAST_TIME_WORDS = [
    ('yesterday', '昨天', '一般过去时标志词'),
    ('last night', '昨晚', '一般过去时标志词'),
    ('last week', '上周', '一般过去时标志词'),
    ('last month', '上个月', '一般过去时标志词'),
    ('last year', '去年', '一般过去时标志词'),
    ('two days ago', '两天前', '一般过去时标志词'),
    ('three years ago', '三年前', '一般过去时标志词'),
    ('just now', '刚才', '一般过去时标志词'),
    ('in 2020', '在2020年', '一般过去时标志词'),
    ('the day before yesterday', '前天', '一般过去时标志词'),
    ('this morning', '今天早上（已过去）', '一般过去时标志词'),
    ('long ago', '很久以前', '一般过去时标志词'),
]

# 动词过去式规则变化
DAY6_PAST_REGULAR = [
    ('play', 'played', '直接加 ed'),
    ('watch', 'watched', '直接加 ed'),
    ('clean', 'cleaned', '直接加 ed'),
    ('visit', 'visited', '直接加 ed'),
    ('help', 'helped', '直接加 ed'),
    ('like', 'liked', '以 e 结尾加 d'),
    ('live', 'lived', '以 e 结尾加 d'),
    ('dance', 'danced', '以 e 结尾加 d'),
    ('hope', 'hoped', '以 e 结尾加 d'),
    ('use', 'used', '以 e 结尾加 d'),
    ('stop', 'stopped', '重读闭音节双写加 ed'),
    ('plan', 'planned', '重读闭音节双写加 ed'),
    ('drop', 'dropped', '重读闭音节双写加 ed'),
    ('study', 'studied', '辅音+y 结尾，变 y 为 i 加 ed'),
    ('carry', 'carried', '辅音+y 结尾，变 y 为 i 加 ed'),
    ('cry', 'cried', '辅音+y 结尾，变 y 为 i 加 ed'),
    ('try', 'tried', '辅音+y 结尾，变 y 为 i 加 ed'),
    ('worry', 'worried', '辅音+y 结尾，变 y 为 i 加 ed'),
    ('play', 'played', '元音+y 结尾，直接加 ed'),
    ('stay', 'stayed', '元音+y 结尾，直接加 ed'),
]

# 动词过去式不规则变化
DAY6_PAST_IRREGULAR = [
    ('is/am', 'was', '不规则变化'),
    ('are', 'were', '不规则变化'),
    ('do', 'did', '不规则变化'),
    ('have/has', 'had', '不规则变化'),
    ('go', 'went', '不规则变化'),
    ('come', 'came', '不规则变化'),
    ('see', 'saw', '不规则变化'),
    ('eat', 'ate', '不规则变化'),
    ('drink', 'drank', '不规则变化'),
    ('give', 'gave', '不规则变化'),
    ('take', 'took', '不规则变化'),
    ('make', 'made', '不规则变化'),
    ('get', 'got', '不规则变化'),
    ('read', 'read', '不规则变化（读音变化）'),
    ('write', 'wrote', '不规则变化'),
    ('run', 'ran', '不规则变化'),
    ('swim', 'swam', '不规则变化'),
    ('sing', 'sang', '不规则变化'),
    ('sit', 'sat', '不规则变化'),
    ('buy', 'bought', '不规则变化'),
    ('bring', 'brought', '不规则变化'),
    ('think', 'thought', '不规则变化'),
    ('teach', 'taught', '不规则变化'),
    ('catch', 'caught', '不规则变化'),
    ('say', 'said', '不规则变化'),
    ('tell', 'told', '不规则变化'),
    ('find', 'found', '不规则变化'),
    ('know', 'knew', '不规则变化'),
    ('fly', 'flew', '不规则变化'),
    ('draw', 'drew', '不规则变化'),
]

# 一般将来时标志词
DAY6_FUTURE_TIME_WORDS = [
    ('tomorrow', '明天', '一般将来时标志词'),
    ('next week', '下周', '一般将来时标志词'),
    ('next month', '下个月', '一般将来时标志词'),
    ('next year', '明年', '一般将来时标志词'),
    ('in two days', '两天后', '一般将来时标志词'),
    ('soon', '很快', '一般将来时标志词'),
    ('this afternoon', '今天下午（未到）', '一般将来时标志词'),
    ('this evening', '今天晚上（未到）', '一般将来时标志词'),
    ('the day after tomorrow', '后天', '一般将来时标志词'),
    ('in the future', '在将来', '一般将来时标志词'),
]

# will 句型练习
DAY6_WILL_SENTENCES = [
    ('I will go to Beijing tomorrow.', '我明天将去北京。', 'will + 动词原形'),
    ('She will visit her grandma next week.', '她下周将去看望奶奶。', 'will + 动词原形'),
    ('They will have a party tonight.', '他们今晚将举办派对。', 'will + 动词原形'),
    ('Will you come to my birthday party?', '你会来我的生日派对吗？', 'Will + 主语 + 动词原形'),
    ('He will not (won\'t) be late.', '他不会迟到。', 'will not/won\'t + 动词原形'),
    ('It will rain tomorrow.', '明天会下雨。', 'will + 动词原形'),
    ('We will finish the work soon.', '我们很快会完成工作。', 'will + 动词原形'),
    ('Will she help us?', '她会帮助我们吗？', 'Will + 主语 + 动词原形'),
    ('I won\'t tell anyone.', '我不会告诉任何人。', 'won\'t + 动词原形'),
    ('What will you do tomorrow?', '你明天要做什么？', '疑问词 + will + 主语 + 动词原形'),
]

# be going to 句型练习
DAY6_BE_GOING_TO = [
    ('I am going to visit my uncle.', '我打算去看望我叔叔。', 'am going to + 动词原形'),
    ('She is going to buy a new dress.', '她打算买一条新裙子。', 'is going to + 动词原形'),
    ('They are going to play football.', '他们打算踢足球。', 'are going to + 动词原形'),
    ('Are you going to have a picnic?', '你们打算去野餐吗？', 'Are + 主语 + going to + 动词原形'),
    ('He is not going to come.', '他不打算来。', 'is not going to + 动词原形'),
    ('Look at the clouds! It is going to rain.', '看那些云！要下雨了。', '根据迹象判断'),
    ('We are going to have a test tomorrow.', '我们明天要考试。', 'are going to + 动词原形'),
    ('What are you going to do this weekend?', '这周末你打算做什么？', '疑问词 + are + 主语 + going to'),
    ('Is she going to join us?', '她打算加入我们吗？', 'Is + 主语 + going to + 动词原形'),
    ('I\'m not going to eat junk food.', '我不打算吃垃圾食品。', 'am not going to + 动词原形'),
]

# 三大时态对比
DAY6_TENSE_COMPARE = [
    ('一般现在时', 'He plays football every day.', '经常性/习惯性动作', 'every day, always, usually'),
    ('一般过去时', 'He played football yesterday.', '过去发生的动作', 'yesterday, last..., ago'),
    ('一般将来时', 'He will play football tomorrow.', '将来要发生的动作', 'tomorrow, next..., will'),
]

# 时态辨析选择题
DAY6_TENSE_CHOOSE = [
    ('He _______ to school yesterday.', 'A. go  B. goes  C. went  D. will go', 'C', 'yesterday 是过去时标志词'),
    ('She _______ TV every evening.', 'A. watch  B. watches  C. watched  D. will watch', 'B', 'every evening 是一般现在时标志词'),
    ('They _______ a picnic next Sunday.', 'A. have  B. has  C. had  D. will have', 'D', 'next Sunday 是将来时标志词'),
    ('I _______ my homework last night.', 'A. do  B. does  C. did  D. will do', 'C', 'last night 是过去时标志词'),
    ('Look! The bus _______. Let\'s run!', 'A. comes  B. came  C. is coming  D. will come', 'C', 'Look! 提示正在发生'),
    ('We _______ to Beijing two years ago.', 'A. go  B. went  C. goes  D. will go', 'B', 'two years ago 是过去时标志词'),
    ('My mother _______ dinner every day.', 'A. cook  B. cooks  C. cooked  D. will cook', 'B', 'every day 是一般现在时标志词'),
    ('Tom _______ his grandpa tomorrow.', 'A. visits  B. visited  C. visit  D. will visit', 'D', 'tomorrow 是将来时标志词'),
    ('She _______ a letter to her friend yesterday.', 'A. write  B. writes  C. wrote  D. will write', 'C', 'yesterday 是过去时标志词'),
    ('_______ you _______ to the party last night?', 'A. Do, go  B. Did, go  C. Will, go  D. Does, go', 'B', 'last night 是过去时标志词'),
    ('It often _______ in summer here.', 'A. rain  B. rains  C. rained  D. will rain', 'B', 'often 是一般现在时标志词'),
    ('I _______ a new bike next month.', 'A. buy  B. bought  C. will buy  D. buys', 'C', 'next month 是将来时标志词'),
    ('What _______ you _______ just now?', 'A. do, do  B. did, do  C. will, do  D. does, do', 'B', 'just now 是过去时标志词'),
    ('She _______ to music every morning.', 'A. listen  B. listens  C. listened  D. will listen', 'B', 'every morning 是一般现在时标志词'),
    ('They _______ a football match next week.', 'A. watch  B. watched  C. watches  D. will watch', 'D', 'next week 是将来时标志词'),
]

# 用所给词适当形式填空
DAY6_FILL_TENSE = [
    ('He _______ (go) to school by bus every day.', 'goes', '一般现在时，第三人称单数'),
    ('She _______ (visit) her grandma last Sunday.', 'visited', '一般过去时，last Sunday'),
    ('They _______ (have) a meeting tomorrow.', 'will have', '一般将来时，tomorrow'),
    ('I _______ (watch) a movie yesterday evening.', 'watched', '一般过去时，yesterday'),
    ('My father _______ (work) in a hospital.', 'works', '一般现在时，描述职业'),
    ('We _______ (go) camping next weekend.', 'will go / are going to go', '一般将来时，next weekend'),
    ('Tom _______ (buy) a new book two days ago.', 'bought', '一般过去时，ago'),
    ('She usually _______ (get) up at 6:30.', 'gets', '一般现在时，usually'),
    ('_______ you _______ (see) the film last night?', 'Did, see', '一般过去时疑问句'),
    ('It _______ (rain) heavily yesterday.', 'rained', '一般过去时，yesterday'),
    ('He _______ (not go) to school tomorrow.', 'won\'t go / is not going to go', '一般将来时否定句'),
    ('My mother _______ (cook) dinner every evening.', 'cooks', '一般现在时，every evening'),
    ('They _______ (play) basketball last week.', 'played', '一般过去时，last week'),
    ('I _______ (be) a teacher in the future.', 'will be', '一般将来时，in the future'),
    ('She _______ (not watch) TV last night.', 'didn\'t watch', '一般过去时否定句'),
    ('_______ he _______ (come) to the party tomorrow?', 'Will, come', '一般将来时疑问句'),
    ('We _______ (have) a good time yesterday.', 'had', '一般过去时，yesterday'),
    ('The sun _______ (rise) in the east.', 'rises', '一般现在时，客观事实'),
    ('I _______ (meet) my friend at the airport next Monday.', 'will meet / am going to meet', '一般将来时'),
    ('She _______ (give) me a present last birthday.', 'gave', '一般过去时，last birthday'),
]

# 时态错题整理（易混点）
DAY6_COMMON_ERRORS = [
    ('He goed to school yesterday.', 'goed → went', 'go 是不规则动词，过去式是 went'),
    ('She will goes to Beijing.', 'will goes → will go', 'will 后接动词原形'),
    ('I am go to visit my uncle.', 'am go → am going', 'be going to 结构'),
    ('They was happy yesterday.', 'was → were', 'They 用 were'),
    ('He didn\'t went to school.', 'went → go', 'didn\'t 后接动词原形'),
    ('She is going to visits her friend.', 'visits → visit', 'be going to 后接动词原形'),
    ('I buyed a new book yesterday.', 'buyed → bought', 'buy 是不规则动词，过去式是 bought'),
    ('He will coming tomorrow.', 'will coming → will come', 'will 后接动词原形'),
    ('Did she went to the park?', 'went → go', 'Did 后接动词原形'),
    ('They are going to played football.', 'played → play', 'be going to 后接动词原形'),
    ('She telled me a story.', 'telled → told', 'tell 是不规则动词，过去式是 told'),
    ('He catched a fish yesterday.', 'catched → caught', 'catch 是不规则动词，过去式是 caught'),
    ('I writed a letter last night.', 'writed → wrote', 'write 是不规则动词，过去式是 wrote'),
    ('She swimmed in the pool.', 'swimmed → swam', 'swim 是不规则动词，过去式是 swam'),
    ('They singed a song yesterday.', 'singed → sang', 'sing 是不规则动词，过去式是 sang'),
]

# 句子补全（三大时态综合）
DAY6_COMPLETE_TENSE = [
    ('他昨天去了北京。', 'He _______ to Beijing yesterday.', 'went'),
    ('她每天早上六点起床。', 'She _______ up at 6:00 every morning.', 'gets'),
    ('我们明天将去野餐。', 'We _______ _______ a picnic tomorrow.', 'will have / are going to have'),
    ('汤姆上周买了一本新书。', 'Tom _______ a new book last week.', 'bought'),
    ('她经常帮助她妈妈。', 'She often _______ her mother.', 'helps'),
    ('他们下周将参观博物馆。', 'They _______ _______ the museum next week.', 'will visit / are going to visit'),
    ('我昨晚看了一部电影。', 'I _______ a movie last night.', 'watched'),
    ('太阳从东方升起。', 'The sun _______ in the east.', 'rises'),
    ('她明天不会来。', 'She _______ _______ tomorrow.', 'won\'t come / isn\'t going to come'),
    ('他们昨天踢足球了。', 'They _______ football yesterday.', 'played'),
    ('我爸爸在医院工作。', 'My father _______ in a hospital.', 'works'),
    ('你明天打算做什么？', 'What _______ you _______ _______ _______ tomorrow?', 'are, going, to, do'),
    ('她上周日去看望了奶奶。', 'She _______ her grandma last Sunday.', 'visited'),
    ('鸟儿在天空飞翔。', 'Birds _______ in the sky.', 'fly'),
    ('我们下个月将有一次考试。', 'We _______ _______ a test next month.', 'will have'),
]


# ============ 第7天：语法专项（介词+代词+名词单复数）题库 ============

# 介词 in/on/at/by 用法
DAY7_PREPOSITION_IN = [
    ('in the morning', '在早上', '在一天中的某段时间'),
    ('in the afternoon', '在下午', '在一天中的某段时间'),
    ('in the evening', '在晚上', '在一天中的某段时间'),
    ('in January', '在一月', '在月份前'),
    ('in 2024', '在2024年', '在年份前'),
    ('in spring', '在春天', '在季节前'),
    ('in summer', '在夏天', '在季节前'),
    ('in China', '在中国', '在国家/城市前'),
    ('in Beijing', '在北京', '在国家/城市前'),
    ('in the classroom', '在教室里', '在较大空间内'),
    ('in the box', '在盒子里', '在封闭空间内'),
    ('in English', '用英语', '用某种语言'),
]

DAY7_PREPOSITION_ON = [
    ('on Monday', '在周一', '在星期几前'),
    ('on Sunday morning', '在周日早上', '在具体某天的某段时间'),
    ('on May 1st', '在五月一日', '在具体日期前'),
    ('on Children\'s Day', '在儿童节', '在节日前'),
    ('on the desk', '在桌子上', '在表面上'),
    ('on the wall', '在墙上', '在表面上'),
    ('on the left', '在左边', '在方位'),
    ('on the right', '在右边', '在方位'),
    ('on foot', '步行', '固定搭配'),
    ('on TV', '在电视上', '固定搭配'),
    ('on the phone', '在打电话', '固定搭配'),
    ('on duty', '值日', '固定搭配'),
]

DAY7_PREPOSITION_AT = [
    ('at 7 o\'clock', '在7点', '在具体时刻前'),
    ('at noon', '在中午', '在中午/夜间'),
    ('at night', '在夜间', '在中午/夜间'),
    ('at the weekend', '在周末', '在周末'),
    ('at home', '在家', '在较小地点'),
    ('at school', '在学校', '在较小地点'),
    ('at the bus stop', '在公交站', '在较小地点'),
    ('at the door', '在门口', '在某一点'),
    ('at the age of 10', '在10岁时', '在某年龄'),
    ('at first', '起初', '固定搭配'),
    ('at last', '最后', '固定搭配'),
    ('at once', '立刻', '固定搭配'),
]

DAY7_PREPOSITION_BY = [
    ('by bus', '乘公交车', '乘坐交通工具'),
    ('by car', '乘汽车', '乘坐交通工具'),
    ('by bike', '骑自行车', '乘坐交通工具'),
    ('by plane', '乘飞机', '乘坐交通工具'),
    ('by train', '乘火车', '乘坐交通工具'),
    ('by the window', '在窗户旁边', '在...旁边'),
    ('by the river', '在河边', '在...旁边'),
    ('by hand', '用手工', '通过某种方式'),
    ('by oneself', '独自', '固定搭配'),
    ('by the way', '顺便说一下', '固定搭配'),
]

# 人称代词（主格/宾格）
DAY7_PERSONAL_PRONOUNS = [
    ('I', 'me', '我', '第一人称单数'),
    ('you', 'you', '你/你们', '第二人称'),
    ('he', 'him', '他', '第三人称单数男'),
    ('she', 'her', '她', '第三人称单数女'),
    ('it', 'it', '它', '第三人称单数物'),
    ('we', 'us', '我们', '第一人称复数'),
    ('they', 'them', '他们/她们/它们', '第三人称复数'),
]

# 人称代词填空练习
DAY7_PRONOUN_FILL = [
    ('_______ am a student. (I/me)', 'I', '主语用主格'),
    ('Please give _______ a book. (I/me)', 'me', '动词后用宾格'),
    ('_______ is my friend. (He/Him)', 'He', '主语用主格'),
    ('I like _______. (he/him)', 'him', '动词后用宾格'),
    ('_______ are good students. (We/Us)', 'We', '主语用主格'),
    ('The teacher teaches _______. (we/us)', 'us', '动词后用宾格'),
    ('_______ is a beautiful girl. (She/Her)', 'She', '主语用主格'),
    ('I know _______. (she/her)', 'her', '动词后用宾格'),
    ('_______ are playing football. (They/Them)', 'They', '主语用主格'),
    ('Please help _______. (they/them)', 'them', '动词后用宾格'),
    ('_______ is a cat. (It/It)', 'It', '主语用主格'),
    ('I feed _______ every day. (it/it)', 'it', '动词后用宾格'),
    ('Let _______ go. (I/me)', 'me', 'let 后用宾格'),
    ('Between you and _______, it\'s a secret. (I/me)', 'me', '介词后用宾格'),
    ('_______ and I are friends. (He/Him)', 'He', '并列主语用主格'),
]

# 物主代词（形容词性/名词性）
DAY7_POSSESSIVE_PRONOUNS = [
    ('my', 'mine', '我的', '第一人称单数'),
    ('your', 'yours', '你的/你们的', '第二人称'),
    ('his', 'his', '他的', '第三人称单数男'),
    ('her', 'hers', '她的', '第三人称单数女'),
    ('its', 'its', '它的', '第三人称单数物'),
    ('our', 'ours', '我们的', '第一人称复数'),
    ('their', 'theirs', '他们的', '第三人称复数'),
]

# 物主代词填空练习
DAY7_POSSESSIVE_FILL = [
    ('This is _______ book. (my/mine)', 'my', '形容词性物主代词+名词'),
    ('The book is _______. (my/mine)', 'mine', '名词性物主代词单独使用'),
    ('_______ name is Tom. (His/He)', 'His', '形容词性物主代词+名词'),
    ('This pen is _______. (her/hers)', 'hers', '名词性物主代词单独使用'),
    ('_______ classroom is big. (Our/Ours)', 'Our', '形容词性物主代词+名词'),
    ('This classroom is _______. (our/ours)', 'ours', '名词性物主代词单独使用'),
    ('_______ dog is cute. (Their/Theirs)', 'Their', '形容词性物主代词+名词'),
    ('The dog is _______. (their/theirs)', 'theirs', '名词性物主代词单独使用'),
    ('Is this _______ bag? (your/yours)', 'your', '形容词性物主代词+名词'),
    ('Is this bag _______? (your/yours)', 'yours', '名词性物主代词单独使用'),
    ('The cat is washing _______ face. (it/its)', 'its', '形容词性物主代词+名词'),
    ('_______ mother is a teacher. (She/Her)', 'Her', '形容词性物主代词+名词'),
    ('This is my book. Where is _______? (your/yours)', 'yours', '名词性物主代词代替 your book'),
    ('My pen is red. _______ is blue. (Her/Hers)', 'Hers', '名词性物主代词代替 her pen'),
    ('Our school is big. _______ is small. (Their/Theirs)', 'Theirs', '名词性物主代词代替 their school'),
]

# 名词单复数规则变化
DAY7_NOUN_PLURAL_REGULAR = [
    ('book', 'books', '直接加 s'),
    ('pen', 'pens', '直接加 s'),
    ('desk', 'desks', '直接加 s'),
    ('cat', 'cats', '直接加 s'),
    ('dog', 'dogs', '直接加 s'),
    ('bus', 'buses', '以 s 结尾加 es'),
    ('box', 'boxes', '以 x 结尾加 es'),
    ('watch', 'watches', '以 ch 结尾加 es'),
    ('brush', 'brushes', '以 sh 结尾加 es'),
    ('class', 'classes', '以 s 结尾加 es'),
    ('tomato', 'tomatoes', '以 o 结尾加 es（有生命）'),
    ('potato', 'potatoes', '以 o 结尾加 es（有生命）'),
    ('photo', 'photos', '以 o 结尾加 s（无生命）'),
    ('piano', 'pianos', '以 o 结尾加 s（无生命）'),
    ('zoo', 'zoos', '以 o 结尾加 s'),
    ('baby', 'babies', '辅音+y 结尾，变 y 为 i 加 es'),
    ('family', 'families', '辅音+y 结尾，变 y 为 i 加 es'),
    ('city', 'cities', '辅音+y 结尾，变 y 为 i 加 es'),
    ('story', 'stories', '辅音+y 结尾，变 y 为 i 加 es'),
    ('boy', 'boys', '元音+y 结尾，直接加 s'),
    ('day', 'days', '元音+y 结尾，直接加 s'),
    ('toy', 'toys', '元音+y 结尾，直接加 s'),
    ('knife', 'knives', '以 f/fe 结尾，变 f/fe 为 v 加 es'),
    ('wife', 'wives', '以 f/fe 结尾，变 f/fe 为 v 加 es'),
    ('leaf', 'leaves', '以 f/fe 结尾，变 f/fe 为 v 加 es'),
    ('life', 'lives', '以 f/fe 结尾，变 f/fe 为 v 加 es'),
    ('half', 'halves', '以 f/fe 结尾，变 f/fe 为 v 加 es'),
]

# 名词单复数不规则变化
DAY7_NOUN_PLURAL_IRREGULAR = [
    ('man', 'men', '不规则变化'),
    ('woman', 'women', '不规则变化'),
    ('child', 'children', '不规则变化'),
    ('foot', 'feet', '不规则变化'),
    ('tooth', 'teeth', '不规则变化'),
    ('mouse', 'mice', '不规则变化'),
    ('goose', 'geese', '不规则变化'),
    ('sheep', 'sheep', '单复数同形'),
    ('deer', 'deer', '单复数同形'),
    ('fish', 'fish', '单复数同形'),
    ('Chinese', 'Chinese', '单复数同形'),
    ('Japanese', 'Japanese', '单复数同形'),
]

# in/on/at 时间用法辨析
DAY7_TIME_PREPOSITION = [
    ('_______ the morning', 'in', '在一天中的某段时间用 in'),
    ('_______ Monday', 'on', '在星期几用 on'),
    ('_______ 8 o\'clock', 'at', '在具体时刻用 at'),
    ('_______ May', 'in', '在月份用 in'),
    ('_______ May 1st', 'on', '在具体日期用 on'),
    ('_______ 2024', 'in', '在年份用 in'),
    ('_______ night', 'at', '在夜间用 at'),
    ('_______ Sunday morning', 'on', '在具体某天的某段时间用 on'),
    ('_______ noon', 'at', '在中午用 at'),
    ('_______ spring', 'in', '在季节用 in'),
    ('_______ the weekend', 'at', '在周末用 at'),
    ('_______ Children\'s Day', 'on', '在节日用 on'),
    ('_______ the evening', 'in', '在一天中的某段时间用 in'),
    ('_______ Friday afternoon', 'on', '在具体某天的某段时间用 on'),
    ('_______ half past seven', 'at', '在具体时刻用 at'),
]

# this/that/these/those 辨析
DAY7_DEMONSTRATIVE = [
    ('_______ is my book. (这是我的书)', 'This', '单数近指用 this'),
    ('_______ are my books. (这些是我的书)', 'These', '复数近指用 these'),
    ('_______ is his pen. (那是他的钢笔)', 'That', '单数远指用 that'),
    ('_______ are his pens. (那些是他的钢笔)', 'Those', '复数远指用 those'),
    ('What\'s _______? (这是什么？)', 'this', '单数近指用 this'),
    ('What are _______? (那些是什么？)', 'those', '复数远指用 those'),
    ('_______ apples are red. (这些苹果是红色的)', 'These', '复数近指用 these'),
    ('_______ apple is green. (那个苹果是绿色的)', 'That', '单数远指用 that'),
    ('Is _______ your bag? (这是你的包吗？)', 'this', '单数近指用 this'),
    ('Are _______ your shoes? (那些是你的鞋子吗？)', 'those', '复数远指用 those'),
    ('_______ is Tom speaking. (我是汤姆，电话用语)', 'This', '电话中介绍自己用 this'),
    ('Is _______ Mary? (你是玛丽吗？电话用语)', 'that', '电话中询问对方用 that'),
]

# 代词使用注意事项（易错题）
DAY7_PRONOUN_ERRORS = [
    ('Me am a student.', 'Me → I', '主语用主格 I'),
    ('Him is my friend.', 'Him → He', '主语用主格 He'),
    ('I like she.', 'she → her', '动词后用宾格 her'),
    ('This is mine book.', 'mine → my', '名词前用形容词性物主代词'),
    ('The book is my.', 'my → mine', '单独使用用名词性物主代词'),
    ('Her is a teacher.', 'Her → She', '主语用主格 She'),
    ('Give I a pen.', 'I → me', '动词后用宾格 me'),
    ('Us are students.', 'Us → We', '主语用主格 We'),
    ('This is her. (这是她的书)', 'her → hers / her book', '单独使用用名词性物主代词'),
    ('Me and Tom are friends.', 'Me → Tom and I', '并列主语用主格，且把 I 放后面'),
    ('Between you and I', 'I → me', '介词后用宾格'),
    ('Let he go.', 'he → him', 'let 后用宾格'),
]

# 介词填空综合练习
DAY7_PREPOSITION_FILL = [
    ('I get up _______ 6:30 every morning.', 'at', '在具体时刻用 at'),
    ('We have English class _______ Monday.', 'on', '在星期几用 on'),
    ('She was born _______ 2012.', 'in', '在年份用 in'),
    ('The book is _______ the desk.', 'on', '在表面上用 on'),
    ('There is a bird _______ the tree.', 'in', '在树上（树内）用 in'),
    ('I go to school _______ bus.', 'by', '乘坐交通工具用 by'),
    ('He lives _______ Beijing.', 'in', '在城市用 in'),
    ('We have lunch _______ noon.', 'at', '在中午用 at'),
    ('The picture is _______ the wall.', 'on', '在墙上用 on'),
    ('I often play football _______ the afternoon.', 'in', '在下午用 in'),
    ('She sits _______ the window.', 'by', '在...旁边用 by'),
    ('We don\'t have class _______ Sunday.', 'on', '在星期几用 on'),
    ('He is good _______ English.', 'at', '擅长用 be good at'),
    ('The cat is _______ the box.', 'in', '在盒子里用 in'),
    ('I\'ll see you _______ 3 o\'clock.', 'at', '在具体时刻用 at'),
]

# 代词替换练习
DAY7_PRONOUN_REPLACE = [
    ('Tom is a student.', 'He is a student.', 'Tom → He'),
    ('Mary likes apples.', 'She likes apples.', 'Mary → She'),
    ('The book is on the desk.', 'It is on the desk.', 'The book → It'),
    ('Tom and I are friends.', 'We are friends.', 'Tom and I → We'),
    ('Give the pen to Tom.', 'Give the pen to him.', 'Tom → him'),
    ('I like Mary.', 'I like her.', 'Mary → her'),
    ('This is Tom\'s book.', 'This is his book.', 'Tom\'s → his'),
    ('The bag is Mary\'s.', 'The bag is hers.', 'Mary\'s → hers'),
    ('Tom and Mary are students.', 'They are students.', 'Tom and Mary → They'),
    ('Please help Tom and Mary.', 'Please help them.', 'Tom and Mary → them'),
    ('This is Tom and Mary\'s classroom.', 'This is their classroom.', 'Tom and Mary\'s → their'),
    ('The classroom is Tom and Mary\'s.', 'The classroom is theirs.', 'Tom and Mary\'s → theirs'),
]

# 名词单复数转换练习
DAY7_NOUN_CONVERT = [
    ('one book → two _______', 'books', '直接加 s'),
    ('one box → three _______', 'boxes', '以 x 结尾加 es'),
    ('one baby → many _______', 'babies', '辅音+y 变 y 为 i 加 es'),
    ('one man → two _______', 'men', '不规则变化'),
    ('one child → three _______', 'children', '不规则变化'),
    ('one foot → two _______', 'feet', '不规则变化'),
    ('one knife → two _______', 'knives', '以 fe 结尾变 fe 为 v 加 es'),
    ('one tomato → many _______', 'tomatoes', '以 o 结尾加 es'),
    ('one photo → two _______', 'photos', '以 o 结尾加 s'),
    ('one sheep → many _______', 'sheep', '单复数同形'),
    ('one tooth → two _______', 'teeth', '不规则变化'),
    ('one family → two _______', 'families', '辅音+y 变 y 为 i 加 es'),
    ('one watch → three _______', 'watches', '以 ch 结尾加 es'),
    ('one leaf → many _______', 'leaves', '以 f 结尾变 f 为 v 加 es'),
    ('one mouse → two _______', 'mice', '不规则变化'),
]


# ============ 第8天：阅读专项（完形填空+阅读理解基础）题库 ============

# 阅读技巧知识点
DAY8_READING_TIPS = [
    ('先看题干再读文章', '带着问题阅读，提高效率', '阅读技巧'),
    ('圈画关键词', '时间、人物、数字、地点是重点', '阅读技巧'),
    ('根据上下文猜词义', '利用前后句理解生词', '阅读技巧'),
    ('注意转折词', 'but, however, although 后常有答案', '阅读技巧'),
    ('首尾句很重要', '段落首句和尾句常包含主旨', '阅读技巧'),
    ('完形填空看搭配', '注意固定短语和习惯用法', '完形技巧'),
    ('阅读理解看原文定位', '答案一定在文章中能找到依据', '阅读技巧'),
]

# 完形填空文章1：My School Day
DAY8_CLOZE_1 = {
    'title': 'My School Day',
    'passage': '''I am a student. I get up at 6:30 every morning. After I wash my face and brush my (1)_______, I have breakfast. I usually have bread and milk (2)_______ breakfast. Then I go to school (3)_______ bike. It takes me about 15 minutes.

I have four classes in the morning. My (4)_______ subject is English because it is very interesting. I have lunch at school. In the afternoon, I have two classes. After school, I often play (5)_______ with my friends. We have a good time.

I go home at 5:00. I do my homework (6)_______ dinner. Then I watch TV for half an hour. I go to bed at 9:00. I have a (7)_______ day every day.''',
    'questions': [
        ('(1)', 'A. face  B. teeth  C. hands  D. feet', 'B', 'brush teeth 刷牙是固定搭配'),
        ('(2)', 'A. for  B. at  C. in  D. on', 'A', 'for breakfast 表示"作为早餐"'),
        ('(3)', 'A. on  B. by  C. in  D. at', 'B', 'by bike 骑自行车'),
        ('(4)', 'A. good  B. bad  C. favourite  D. difficult', 'C', '根据后文 interesting 可知是最喜欢的'),
        ('(5)', 'A. basketball  B. the basketball  C. a basketball  D. basketballs', 'A', 'play basketball 不加冠词'),
        ('(6)', 'A. after  B. before  C. in  D. at', 'A', '根据常理，先做作业再吃晚饭或吃完晚饭做作业'),
        ('(7)', 'A. sad  B. bad  C. busy  D. free', 'C', '根据全文描述，一天很忙碌'),
    ]
}

# 完形填空文章2：My Best Friend
DAY8_CLOZE_2 = {
    'title': 'My Best Friend',
    'passage': '''I have a good friend. (1)_______ name is Tom. He is 12 years old. He is (2)_______ America. He has blue eyes and yellow hair. He is tall and (3)_______.

Tom and I are in the same class. We often help (4)_______. He is good at math, and I am good at English. We (5)_______ from each other.

After school, we like to play football together. (6)_______ weekends, we sometimes go to the park. We have a lot of (7)_______ together. Tom is my best friend.''',
    'questions': [
        ('(1)', 'A. He  B. His  C. Him  D. Her', 'B', '名词前用形容词性物主代词 His'),
        ('(2)', 'A. in  B. at  C. from  D. to', 'C', 'be from 来自'),
        ('(3)', 'A. fat  B. thin  C. short  D. strong', 'D', '根据上下文，tall and strong 搭配合理'),
        ('(4)', 'A. other  B. others  C. each other  D. another', 'C', 'help each other 互相帮助'),
        ('(5)', 'A. learn  B. play  C. come  D. go', 'A', 'learn from 向...学习'),
        ('(6)', 'A. In  B. On  C. At  D. For', 'B', 'On weekends 在周末'),
        ('(7)', 'A. time  B. fun  C. work  D. food', 'B', 'have fun 玩得开心'),
    ]
}

# 阅读理解文章1：Tom's Family
DAY8_READING_1 = {
    'title': 'Tom\'s Family',
    'passage': '''Tom is an American boy. He is twelve years old. He lives in Beijing with his family now.

There are four people in his family: his father, his mother, his sister and him. His father is a doctor. He works in a hospital. His mother is a teacher. She teaches English in a middle school. His sister, Mary, is only five years old. She doesn\'t go to school. She stays at home every day.

Tom studies in a primary school. He likes his school very much. He has many friends there. He often plays football with them after school. He can speak a little Chinese now. He thinks Chinese is very interesting.''',
    'questions': [
        ('1. How old is Tom?', 'A. 10  B. 11  C. 12  D. 13', 'C', '文中说 He is twelve years old'),
        ('2. Where does Tom live now?', 'A. In America  B. In Beijing  C. In Shanghai  D. In London', 'B', '文中说 He lives in Beijing'),
        ('3. What does Tom\'s father do?', 'A. A teacher  B. A doctor  C. A worker  D. A driver', 'B', '文中说 His father is a doctor'),
        ('4. How many people are there in Tom\'s family?', 'A. 3  B. 4  C. 5  D. 6', 'B', '文中说 There are four people'),
        ('5. What does Tom think of Chinese?', 'A. Difficult  B. Easy  C. Boring  D. Interesting', 'D', '文中说 He thinks Chinese is very interesting'),
    ]
}

# 阅读理解文章2：A Busy Saturday
DAY8_READING_2 = {
    'title': 'A Busy Saturday',
    'passage': '''Last Saturday was a busy day for Li Ming. In the morning, he got up at 7:00. After breakfast, he helped his mother clean the house. It took him two hours.

At 11:00, he went to the library with his friend Wang Wei. They read books there for one hour. Then they had lunch at a restaurant near the library.

In the afternoon, Li Ming went to the park. He played basketball with some boys there. He was very happy. At 5:00, he went home. In the evening, he did his homework and watched TV. He went to bed at 9:30.

Li Ming was tired but happy. He had a wonderful Saturday.''',
    'questions': [
        ('1. What did Li Ming do first after breakfast?', 'A. Went to the library  B. Cleaned the house  C. Played basketball  D. Did homework', 'B', '文中说 he helped his mother clean the house'),
        ('2. How long did Li Ming read books in the library?', 'A. One hour  B. Two hours  C. Three hours  D. Half an hour', 'A', '文中说 They read books there for one hour'),
        ('3. Where did Li Ming have lunch?', 'A. At home  B. At school  C. At a restaurant  D. In the park', 'C', '文中说 they had lunch at a restaurant'),
        ('4. What did Li Ming do in the afternoon?', 'A. Read books  B. Cleaned the house  C. Watched TV  D. Played basketball', 'D', '文中说 He played basketball'),
        ('5. How did Li Ming feel about his Saturday?', 'A. Sad  B. Bored  C. Tired but happy  D. Angry', 'C', '文中说 Li Ming was tired but happy'),
    ]
}

# 猜词义练习
DAY8_GUESS_WORD = [
    ('The old man is very kind. He often helps others.', 'kind', '善良的', '根据后文"帮助他人"推断'),
    ('It\'s raining heavily. Don\'t forget to take an umbrella.', 'umbrella', '雨伞', '根据"下大雨"推断'),
    ('The elephant is huge. It\'s the biggest animal in the zoo.', 'huge', '巨大的', '根据"最大的动物"推断'),
    ('She was so tired that she fell asleep quickly.', 'fell asleep', '睡着了', '根据"很累"推断'),
    ('The food is delicious. I want to eat more.', 'delicious', '美味的', '根据"想吃更多"推断'),
    ('He is very clever. He can answer all the questions.', 'clever', '聪明的', '根据"能回答所有问题"推断'),
]


# ============ 第9天：写作专项（小作文·审题+句型+书写）题库 ============

# 审题技巧
DAY9_WRITING_TIPS = [
    ('抓主题', '明确作文要写什么内容', '审题技巧'),
    ('定人称', '确定用第一人称(I)还是第三人称(He/She)', '审题技巧'),
    ('定时态', '根据题目确定用一般现在时、过去时还是将来时', '审题技巧'),
    ('列提纲', '开头-中间-结尾，每部分写什么', '写作技巧'),
    ('字数要求', '小学作文一般要求60-80词', '写作要求'),
]

# 常用开头句
DAY9_OPENING_SENTENCES = [
    ('I want to tell you about...', '我想告诉你关于...', '引出话题'),
    ('I have a...', '我有一个...', '介绍人/物'),
    ('My favourite... is...', '我最喜欢的...是...', '表达喜好'),
    ('I like... very much.', '我非常喜欢...', '表达喜好'),
    ('Let me tell you something about...', '让我告诉你一些关于...的事', '引出话题'),
    ('I\'m going to tell you about...', '我将要告诉你关于...', '引出话题'),
    ('Do you know...?', '你知道...吗？', '引起兴趣'),
    ('I had a wonderful/great/happy...', '我度过了一个美好的...', '描述经历'),
]

# 常用中间句（描述句）
DAY9_BODY_SENTENCES = [
    ('He/She is... years old.', '他/她...岁了。', '描述年龄'),
    ('He/She has... hair and... eyes.', '他/她有...头发和...眼睛。', '描述外貌'),
    ('He/She is tall/short/thin/fat.', '他/她高/矮/瘦/胖。', '描述身材'),
    ('He/She likes...', '他/她喜欢...', '描述爱好'),
    ('He/She is good at...', '他/她擅长...', '描述特长'),
    ('We often... together.', '我们经常一起...', '描述活动'),
    ('In the morning/afternoon/evening, I...', '在早上/下午/晚上，我...', '描述时间活动'),
    ('First... Then... Finally...', '首先...然后...最后...', '描述顺序'),
    ('I think... is very...', '我认为...非常...', '表达观点'),
    ('It was sunny/rainy/cloudy.', '天气晴朗/下雨/多云。', '描述天气'),
]

# 常用结尾句
DAY9_ENDING_SENTENCES = [
    ('What about you?', '你呢？', '询问对方'),
    ('I love him/her very much.', '我非常爱他/她。', '表达感情'),
    ('I had a great time.', '我玩得很开心。', '总结感受'),
    ('I hope you can...', '我希望你能...', '表达希望'),
    ('I\'m looking forward to...', '我期待着...', '表达期待'),
    ('That\'s all. Thank you!', '就这些。谢谢！', '结束语'),
    ('I will never forget it.', '我永远不会忘记它。', '强调印象'),
    ('What a wonderful day!', '多么美好的一天！', '感叹句结尾'),
]

# 写作题目1：My Winter Holiday
DAY9_WRITING_TOPIC_1 = {
    'title': 'My Winter Holiday',
    'requirements': '请以"My Winter Holiday"为题，写一篇60-80词的短文。',
    'hints': ['When was your winter holiday?', 'What did you do?', 'How did you feel?'],
    'outline': {
        'opening': '介绍寒假时间',
        'body': '描述寒假活动（2-3个活动）',
        'ending': '总结感受'
    },
    'sample': '''My Winter Holiday

I had a wonderful winter holiday. It was from January to February.

During the holiday, I did many things. First, I did my homework every day. Then, I visited my grandparents with my parents. We had a big dinner together. I also played with my friends. We made a snowman. It was very fun.

I had a great time. I love my winter holiday very much. What about you?

(About 70 words)'''
}

# 写作题目2：My Best Friend
DAY9_WRITING_TOPIC_2 = {
    'title': 'My Best Friend',
    'requirements': '请以"My Best Friend"为题，写一篇60-80词的短文。',
    'hints': ['Who is your best friend?', 'What does he/she look like?', 'What do you often do together?'],
    'outline': {
        'opening': '介绍好朋友是谁',
        'body': '描述外貌、性格、爱好',
        'ending': '表达感情'
    },
    'sample': '''My Best Friend

I have a best friend. Her name is Li Hua. She is 12 years old.

She has long black hair and big eyes. She is tall and thin. She is very kind and friendly. She likes reading books and playing the piano. She is good at English.

We often do homework together. We also play games after school. She always helps me when I have problems.

I love her very much. She is my best friend forever.

(About 75 words)'''
}

# 写作常见错误
DAY9_WRITING_ERRORS = [
    ('I very like English.', 'I like English very much.', 'very 不能直接修饰动词'),
    ('She have long hair.', 'She has long hair.', '第三人称单数用 has'),
    ('I am go to school.', 'I go to school.', '一般现在时不需要 am'),
    ('He is a my friend.', 'He is my friend.', 'a 和 my 不能同时用'),
    ('I played with he.', 'I played with him.', '介词后用宾格'),
    ('We are have fun.', 'We have fun.', 'have 是实义动词，不需要 are'),
    ('I goed to the park.', 'I went to the park.', 'go 的过去式是 went'),
    ('She is very beautiful girl.', 'She is a very beautiful girl.', '单数可数名词前要加 a'),
]


# ============ 第10天：综合模考+错题复盘+知识梳理 题库 ============

# 综合模考 - 词汇部分
DAY10_VOCAB_TEST = [
    ('look _______', 'for / at / after', '短语搭配', 'look for 寻找, look at 看, look after 照顾'),
    ('get _______', 'up', '短语搭配', 'get up 起床'),
    ('turn _______ the light', 'on / off', '短语搭配', 'turn on 打开, turn off 关闭'),
    ('go to _______', 'school / bed', '短语搭配', 'go to school 上学, go to bed 睡觉'),
    ('play _______', 'football / the piano', '短语搭配', '球类不加 the，乐器加 the'),
    ('_______ the morning', 'in', '介词填空', 'in the morning'),
    ('_______ Monday', 'on', '介词填空', 'on Monday'),
    ('_______ 8 o\'clock', 'at', '介词填空', 'at 8 o\'clock'),
    ('go to school _______ bus', 'by', '介词填空', 'by bus'),
    ('book → _______', 'books', '名词复数', '直接加 s'),
    ('box → _______', 'boxes', '名词复数', '以 x 结尾加 es'),
    ('baby → _______', 'babies', '名词复数', '辅音+y 变 y 为 i 加 es'),
    ('child → _______', 'children', '名词复数', '不规则变化'),
    ('foot → _______', 'feet', '名词复数', '不规则变化'),
    ('sheep → _______', 'sheep', '名词复数', '单复数同形'),
]

# 综合模考 - 语法部分
DAY10_GRAMMAR_TEST = [
    ('He _______ (play) football every day.', 'plays', '一般现在时第三人称单数'),
    ('She _______ (go) to school yesterday.', 'went', '一般过去时'),
    ('They _______ (visit) the museum tomorrow.', 'will visit', '一般将来时'),
    ('_______ you a student? Yes, I _______.', 'Are, am', 'be 动词用法'),
    ('_______ he like apples? Yes, he _______.', 'Does, does', 'do/does 用法'),
    ('He _______ (not go) to school last Sunday.', 'didn\'t go', '过去时否定句'),
    ('She _______ (have) a cat.', 'has', 'have 的第三人称单数'),
    ('_______ is your name? My name is Tom.', 'What', '特殊疑问词'),
    ('_______ do you go to school? By bus.', 'How', '特殊疑问词'),
    ('_______ is your birthday? It\'s May 1st.', 'When', '特殊疑问词'),
    ('This is _______ (I) book.', 'my', '形容词性物主代词'),
    ('The book is _______ (I).', 'mine', '名词性物主代词'),
    ('Please give _______ (he) a pen.', 'him', '人称代词宾格'),
    ('_______ (She) is a teacher.', 'She', '人称代词主格'),
    ('There _______ (be) some water in the glass.', 'is', 'there be 句型'),
    ('How many books _______ (be) there?', 'are', 'there be 句型'),
    ('He is taller _______ me.', 'than', '比较级'),
    ('She is the _______ (tall) in her class.', 'tallest', '最高级'),
    ('I _______ (watch) TV now.', 'am watching', '现在进行时'),
    ('Look! The children _______ (play) games.', 'are playing', '现在进行时'),
]

# 综合模考 - 阅读理解
DAY10_READING_TEST = {
    'title': 'A Letter from Mike',
    'passage': '''Dear Tom,

How are you? I\'m fine. Let me tell you about my school life.

I get up at 6:30 every morning. I have breakfast at 7:00. Then I go to school by bike. School starts at 8:00. I have four classes in the morning and two in the afternoon. My favourite subject is English. I think it\'s very interesting.

I have lunch at school. The food is delicious. After school, I often play basketball with my friends. I go home at 5:00. In the evening, I do my homework and read books.

I like my school life very much. What about you? Please write to me soon.

Yours,
Mike''',
    'questions': [
        ('1. What time does Mike get up?', 'A. 6:00  B. 6:30  C. 7:00  D. 7:30', 'B', '文中说 I get up at 6:30'),
        ('2. How does Mike go to school?', 'A. By bus  B. By car  C. By bike  D. On foot', 'C', '文中说 I go to school by bike'),
        ('3. What is Mike\'s favourite subject?', 'A. Math  B. Chinese  C. English  D. Music', 'C', '文中说 My favourite subject is English'),
        ('4. What does Mike do after school?', 'A. Does homework  B. Plays basketball  C. Reads books  D. Watches TV', 'B', '文中说 I often play basketball'),
        ('5. Where does Mike have lunch?', 'A. At home  B. At school  C. At a restaurant  D. In the park', 'B', '文中说 I have lunch at school'),
    ]
}

# 综合模考 - 写作题目
DAY10_WRITING_TEST = {
    'title': 'My Favourite Season',
    'requirements': '请以"My Favourite Season"为题，写一篇60-80词的短文。',
    'hints': ['What is your favourite season?', 'What is the weather like?', 'What can you do in this season?'],
}

# 知识梳理 - 10天学习要点
DAY10_REVIEW_POINTS = [
    ('第2天', '词汇基础', '短语搭配：look for/at/after, get up, turn on/off 等'),
    ('第3天', '句型基础', '主系表、主谓宾结构，陈述句转疑问句/否定句'),
    ('第4天', '特殊疑问句', 'what/when/where/who/why/how 用法，情景交际'),
    ('第5天', '一般现在时', '动词第三人称单数变化，时间标志词'),
    ('第6天', '一般过去时+将来时', '动词过去式变化，will/be going to'),
    ('第7天', '介词+代词+名词', 'in/on/at/by，人称代词，物主代词，名词单复数'),
    ('第8天', '阅读专项', '完形填空技巧，阅读理解技巧'),
    ('第9天', '写作专项', '审题技巧，常用句型，写作框架'),
    ('第10天', '综合复习', '综合模考，错题复盘，知识梳理'),
]

# 易错点汇总
DAY10_ERROR_SUMMARY = [
    ('动词第三人称单数', 'He plays football. / She has a cat.', '注意 have→has, do→does, go→goes'),
    ('动词过去式', 'He went to school. / She bought a book.', '注意不规则动词变化'),
    ('介词用法', 'in the morning, on Monday, at 8 o\'clock', '时间介词 in/on/at 区分'),
    ('人称代词', 'I-me, he-him, she-her, we-us, they-them', '主格做主语，宾格做宾语'),
    ('物主代词', 'my-mine, your-yours, his-his, her-hers', '形容词性+名词，名词性单独用'),
    ('名词复数', 'boxes, babies, children, feet, sheep', '注意不规则变化和特殊规则'),
    ('there be 句型', 'There is a book. / There are some books.', '就近原则'),
    ('特殊疑问句', 'What is your name? / How do you go to school?', '疑问词+一般疑问句语序'),
]


# ============ 生成第2天文档函数 ============

def generate_day2_doc(version, output_path):
    """生成第2天词汇基础练习题"""
    doc = Document()

    # 根据版本确定题目数量
    if version == '简洁版':
        trans_count, fill_count, dist_count, comp_count = 4, 4, 2, 3
    elif version == '完整版':
        trans_count, fill_count, dist_count, comp_count = 8, 8, 4, 6
    else:  # 充实版
        trans_count, fill_count, dist_count, comp_count = 15, 15, 5, 10

    add_title(doc, '小学六年级英语练习题【词汇基础】', f'（基础+提升）第2天 - {version}')
    add_section_title(doc, '短语与固定搭配专项练习')

    answers = {}

    # 一、短语汉译英
    add_question_header(doc, '一', '短语汉译英（根据中文写出英文短语）')
    ans_list = []
    for i, (cn, en) in enumerate(DAY2_PHRASES_TRANSLATE[:trans_count], 1):
        add_question(doc, f'{i}. {cn} _______________________')
        ans_list.append(f'{i}. {en}')
    answers['一、短语汉译英'] = ans_list

    # 二、选词填空
    add_question_header(doc, '二', '选词填空（从词库中选择正确短语填入句子）')
    word_bank = ['look for', 'look after', 'get up', 'put on', 'turn off',
                 'at night', 'in the morning', 'look at', 'wake up', 'turn on',
                 'take off', 'on time', 'in time', 'go to bed', 'on the weekend']
    add_question(doc, f'词库：{", ".join(word_bank[:min(fill_count+2, len(word_bank))])}')
    doc.add_paragraph()
    ans_list = []
    for i, (sent, ans, _) in enumerate(DAY2_FILL_BLANKS[:fill_count], 1):
        add_question(doc, f'{i}. {sent}')
        ans_list.append(f'{i}. {ans}')
    answers['二、选词填空'] = ans_list

    # 三、短语辨析
    add_question_header(doc, '三', '短语辨析（选择正确的短语填空）')
    ans_list = []
    q_num = 1
    for phrase_group, questions in DAY2_DISTINGUISH[:dist_count]:
        add_question(doc, f'【{phrase_group}】', indent=False)
        for sent, ans in questions:
            add_question(doc, f'{q_num}. {sent}')
            ans_list.append(f'{q_num}. {ans}')
            q_num += 1
    answers['三、短语辨析'] = ans_list

    # 四、完成句子
    add_question_header(doc, '四', '完成句子（根据中文提示完成英文句子）')
    ans_list = []
    for i, (cn, en, ans) in enumerate(DAY2_COMPLETE[:comp_count], 1):
        add_question(doc, f'{i}. {cn}')
        add_question(doc, f'   {en}')
        ans_list.append(f'{i}. {ans}')
    answers['四、完成句子'] = ans_list

    add_answer_section(doc, answers)
    doc.save(output_path)
    print(f'已生成: {output_path}')


# ============ 生成第3天文档函数 ============

def generate_day3_doc(version, output_path):
    """生成第3天句型基础练习题"""
    doc = Document()

    # 根据版本确定题目数量
    if version == '简洁版':
        judge_count, quest_count, neg_count, fill_count, trans_count = 4, 4, 4, 4, 4
    elif version == '完整版':
        judge_count, quest_count, neg_count, fill_count, trans_count = 8, 8, 8, 8, 8
    else:  # 充实版
        judge_count, quest_count, neg_count, fill_count, trans_count = 15, 15, 15, 15, 15

    add_title(doc, '小学六年级英语练习题【句型基础】', f'（基础+提升）第3天 - {version}')
    add_section_title(doc, '陈述句、一般疑问句与否定句专项练习')

    answers = {}

    # 一、句型判断
    add_question_header(doc, '一', '句型判断（判断下列句子是"主系表"还是"主谓宾"结构）')
    ans_list = []
    for i, (sent, ans, reason) in enumerate(DAY3_SENTENCE_JUDGE[:judge_count], 1):
        add_question(doc, f'{i}. {sent}  （        ）')
        ans_list.append(f'{i}. {ans}（{reason}）')
    answers['一、句型判断'] = ans_list

    # 二、陈述句转一般疑问句
    add_question_header(doc, '二', '陈述句转一般疑问句')
    ans_list = []
    for i, (sent, ans, rule) in enumerate(DAY3_TO_QUESTION[:quest_count], 1):
        add_question(doc, f'{i}. {sent}')
        add_question(doc, '   _______________________________________')
        ans_list.append(f'{i}. {ans}（{rule}）')
    answers['二、陈述句转一般疑问句'] = ans_list

    # 三、陈述句转否定句
    add_question_header(doc, '三', '陈述句转否定句')
    ans_list = []
    for i, (sent, ans, rule) in enumerate(DAY3_TO_NEGATIVE[:neg_count], 1):
        add_question(doc, f'{i}. {sent}')
        add_question(doc, '   _______________________________________')
        ans_list.append(f'{i}. {ans}（{rule}）')
    answers['三、陈述句转否定句'] = ans_list

    # 四、be动词与do/does选择填空
    add_question_header(doc, '四', 'be动词与do/does选择填空')
    ans_list = []
    for i, (sent, ans, rule) in enumerate(DAY3_BE_DO_FILL[:fill_count], 1):
        add_question(doc, f'{i}. {sent}')
        ans_list.append(f'{i}. {ans}（{rule}）')
    answers['四、be动词与do/does选择填空'] = ans_list

    # 五、句型转换综合练习
    add_question_header(doc, '五', '句型转换综合练习（按要求改写句子）')
    ans_list = []
    for i, (sent, req, ans) in enumerate(DAY3_TRANSFORM[:trans_count], 1):
        add_question(doc, f'{i}. {sent}（改为{req}）')
        add_question(doc, '   _______________________________________')
        ans_list.append(f'{i}. {ans}')
    answers['五、句型转换综合练习'] = ans_list

    add_answer_section(doc, answers)
    doc.save(output_path)
    print(f'已生成: {output_path}')


# ============ 生成第4天文档函数 ============

def generate_day4_doc(version, output_path):
    """生成第4天特殊疑问句+情景交际练习题"""
    doc = Document()

    # 根据版本确定题目数量
    if version == '简洁版':
        wh_fill, wh_match, wh_order = 4, 4, 4
        shop, way, greet = 2, 2, 2
        make_sent, dialogue_count = 4, 2
    elif version == '完整版':
        wh_fill, wh_match, wh_order = 8, 8, 8
        shop, way, greet = 4, 4, 4
        make_sent, dialogue_count = 6, 3
    else:  # 充实版
        wh_fill, wh_match, wh_order = 15, 12, 15
        shop, way, greet = 8, 8, 10
        make_sent, dialogue_count = 12, 5

    add_title(doc, '小学六年级英语练习题【特殊疑问句+情景交际】', f'（基础+提升）第4天 - {version}')

    answers = {}

    # ========== 第一部分：特殊疑问词 ==========
    add_section_title(doc, '第一部分：特殊疑问词（20分钟）')

    # 一、疑问词选择填空
    add_question_header(doc, '一', '疑问词选择填空（从 what/when/where/who/why/how 中选择）')
    ans_list = []
    for i, (sent, ans, rule) in enumerate(DAY4_WH_FILL[:wh_fill], 1):
        add_question(doc, f'{i}. {sent}')
        ans_list.append(f'{i}. {ans}（{rule}）')
    answers['一、疑问词选择填空'] = ans_list

    # 二、疑问词意义匹配
    add_question_header(doc, '二', '疑问词意义匹配（将疑问词与其含义连线或填写）')
    add_question(doc, '请将左边的疑问词与右边的含义进行匹配：', indent=False)
    doc.add_paragraph()
    ans_list = []
    match_items = DAY4_WH_MATCH[:wh_match]
    for i, (word, meaning, usage) in enumerate(match_items, 1):
        add_question(doc, f'{i}. {word}        ______')
        ans_list.append(f'{i}. {word} — {meaning}（{usage}）')
    doc.add_paragraph()
    add_question(doc, '备选含义：', indent=False)
    meanings = [item[1] for item in match_items]
    import random
    shuffled = meanings.copy()
    random.seed(42)  # 固定随机种子保证一致性
    random.shuffle(shuffled)
    add_question(doc, f'A. {shuffled[0]}  B. {shuffled[1]}  C. {shuffled[2]}  D. {shuffled[3]}' if len(shuffled) >= 4 else ', '.join(shuffled))
    if len(shuffled) > 4:
        add_question(doc, f'E. {shuffled[4]}  F. {shuffled[5]}' + (f'  G. {shuffled[6]}  H. {shuffled[7]}' if len(shuffled) >= 8 else ''))
    answers['二、疑问词意义匹配'] = ans_list

    # 三、特殊疑问句语序判断
    add_question_header(doc, '三', '特殊疑问句语序判断（判断下列句子语序是否正确，错误的请改正）')
    ans_list = []
    for i, (sent, ans, rule) in enumerate(DAY4_WH_ORDER[:wh_order], 1):
        add_question(doc, f'{i}. {sent}  （    ）')
        ans_list.append(f'{i}. {ans}（{rule}）')
    answers['三、特殊疑问句语序判断'] = ans_list

    # ========== 第二部分：情景交际 ==========
    add_section_title(doc, '第二部分：情景交际（15分钟）')

    # 四、购物场景应答
    add_question_header(doc, '四', '购物场景应答（根据对话情景，选择或填写合适的应答）')
    ans_list = []
    for i, (dialogue, ans, tip) in enumerate(DAY4_SHOPPING[:shop], 1):
        add_question(doc, f'{i}. {dialogue}')
        ans_list.append(f'{i}. {ans}（{tip}）')
    answers['四、购物场景应答'] = ans_list

    # 五、问路场景应答
    add_question_header(doc, '五', '问路场景应答（根据对话情景，选择或填写合适的应答）')
    ans_list = []
    for i, (dialogue, ans, tip) in enumerate(DAY4_ASKING_WAY[:way], 1):
        add_question(doc, f'{i}. {dialogue}')
        ans_list.append(f'{i}. {ans}（{tip}）')
    answers['五、问路场景应答'] = ans_list

    # 六、问候场景应答
    add_question_header(doc, '六', '问候场景应答（根据对话情景，选择或填写合适的应答）')
    ans_list = []
    for i, (dialogue, ans, tip) in enumerate(DAY4_GREETING[:greet], 1):
        add_question(doc, f'{i}. {dialogue}')
        ans_list.append(f'{i}. {ans}（{tip}）')
    answers['六、问候场景应答'] = ans_list

    # ========== 第三部分：综合练习 ==========
    add_section_title(doc, '第三部分：综合练习（15分钟）')

    # 七、连词成句
    add_question_header(doc, '七', '连词成句（将打乱的单词组成正确的特殊疑问句）')
    ans_list = []
    for i, (words, ans, tip) in enumerate(DAY4_MAKE_SENTENCE[:make_sent], 1):
        add_question(doc, f'{i}. {words}')
        add_question(doc, '   _______________________________________')
        ans_list.append(f'{i}. {ans}（{tip}）')
    answers['七、连词成句'] = ans_list

    # 八、补全对话
    add_question_header(doc, '八', '补全对话（根据上下文，在横线处填入合适的句子）')
    ans_list = []
    for i, (scene, dlg_text, ans_items) in enumerate(DAY4_DIALOGUE_COMPLETE[:dialogue_count], 1):
        add_question(doc, f'【{scene}】', indent=False)
        for line in dlg_text.split('\n'):
            add_question(doc, line)
        doc.add_paragraph()
        ans_list.append(f'{i}. {scene}：' + ' / '.join([f'({j+1}) {a}' for j, a in enumerate(ans_items)]))
    answers['八、补全对话'] = ans_list

    add_answer_section(doc, answers)
    doc.save(output_path)
    print(f'已生成: {output_path}')


# ============ 生成第5天文档函数 ============

def generate_day5_doc(version, output_path):
    """生成第5天语法专项（一般现在时）练习题"""
    doc = Document()

    # 根据版本确定题目数量
    if version == '简洁版':
        verb_change, tense_judge, verb_choose = 5, 4, 4
        rule_count, time_words, irregular = 2, 4, 3
        fill_blank, correct_error, complete_sent = 4, 4, 4
    elif version == '完整版':
        verb_change, tense_judge, verb_choose = 10, 8, 8
        rule_count, time_words, irregular = 3, 8, 6
        fill_blank, correct_error, complete_sent = 8, 8, 8
    else:  # 充实版
        verb_change, tense_judge, verb_choose = 20, 15, 15
        rule_count, time_words, irregular = 5, 12, 10
        fill_blank, correct_error, complete_sent = 15, 15, 15

    add_title(doc, '小学六年级英语练习题【语法专项：一般现在时】', f'（基础+提升）第5天 - {version}')

    answers = {}

    # ========== 第一部分：语法精讲 ==========
    add_section_title(doc, '第一部分：语法精讲（20分钟）')

    # 一、动词变形练习
    add_question_header(doc, '一', '动词变形练习（写出下列动词的第三人称单数形式）')
    ans_list = []
    for i, (verb, third, rule) in enumerate(DAY5_VERB_CHANGE[:verb_change], 1):
        add_question(doc, f'{i}. {verb} → _______')
        ans_list.append(f'{i}. {third}（{rule}）')
    answers['一、动词变形练习'] = ans_list

    # 二、时态判断
    add_question_header(doc, '二', '时态判断（判断下列句子是否为一般现在时，是写"是"，否写"否"）')
    ans_list = []
    for i, (sent, ans, reason) in enumerate(DAY5_TENSE_JUDGE[:tense_judge], 1):
        add_question(doc, f'{i}. {sent}  （    ）')
        ans_list.append(f'{i}. {ans}（{reason}）')
    answers['二、时态判断'] = ans_list

    # 三、动词形式选择
    add_question_header(doc, '三', '动词形式选择（从括号中选择正确的动词形式）')
    ans_list = []
    for i, (sent, ans, reason) in enumerate(DAY5_VERB_CHOOSE[:verb_choose], 1):
        add_question(doc, f'{i}. {sent}')
        ans_list.append(f'{i}. {ans}（{reason}）')
    answers['三、动词形式选择'] = ans_list

    # 四、变形规则分类
    add_question_header(doc, '四', '变形规则分类（将下列动词按变形规则分类）')
    add_question(doc, '请将动词填入对应的规则类别中：', indent=False)
    doc.add_paragraph()
    # 收集所有动词用于展示
    all_verbs = []
    for rule, verbs in DAY5_RULE_CLASSIFY[:rule_count]:
        all_verbs.extend(verbs[:4])  # 每类取前4个
    add_question(doc, f'动词：{", ".join(all_verbs[:12])}')
    doc.add_paragraph()
    ans_list = []
    for i, (rule, verbs) in enumerate(DAY5_RULE_CLASSIFY[:rule_count], 1):
        add_question(doc, f'{i}. {rule}：_______________________')
        ans_list.append(f'{i}. {rule}：{", ".join(verbs)}')
    answers['四、变形规则分类'] = ans_list

    # ========== 第二部分：答题技巧 ==========
    add_section_title(doc, '第二部分：答题技巧（10分钟）')

    # 五、时间标志词识别
    add_question_header(doc, '五', '时间标志词识别（写出下列标志词的中文意思）')
    ans_list = []
    for i, (word, meaning, tip) in enumerate(DAY5_TIME_WORDS[:time_words], 1):
        add_question(doc, f'{i}. {word} → _______')
        ans_list.append(f'{i}. {meaning}（{tip}）')
    answers['五、时间标志词识别'] = ans_list

    # 六、不规则变化专练
    add_question_header(doc, '六', '不规则变化专练（用括号内动词的正确形式填空）')
    ans_list = []
    for i, (verb, third, sent, ans) in enumerate(DAY5_IRREGULAR[:irregular], 1):
        add_question(doc, f'{i}. {sent}')
        ans_list.append(f'{i}. {ans}（{verb}→{third}）')
    answers['六、不规则变化专练'] = ans_list

    # ========== 第三部分：专项练习 ==========
    add_section_title(doc, '第三部分：专项练习（20分钟）')

    # 七、时态填空
    add_question_header(doc, '七', '时态填空（用括号内动词的正确形式填空）')
    ans_list = []
    for i, (sent, ans, reason) in enumerate(DAY5_FILL_BLANK[:fill_blank], 1):
        add_question(doc, f'{i}. {sent}')
        ans_list.append(f'{i}. {ans}（{reason}）')
    answers['七、时态填空'] = ans_list

    # 八、单句改错
    add_question_header(doc, '八', '单句改错（找出句中错误并改正）')
    ans_list = []
    for i, (sent, correction, reason) in enumerate(DAY5_CORRECT_ERROR[:correct_error], 1):
        add_question(doc, f'{i}. {sent}')
        add_question(doc, '   错误：_______ → 改正：_______')
        ans_list.append(f'{i}. {correction}（{reason}）')
    answers['八、单句改错'] = ans_list

    # 九、句子补全
    add_question_header(doc, '九', '句子补全（根据中文提示完成英文句子）')
    ans_list = []
    for i, (cn, en, ans) in enumerate(DAY5_COMPLETE_SENT[:complete_sent], 1):
        add_question(doc, f'{i}. {cn}')
        add_question(doc, f'   {en}')
        ans_list.append(f'{i}. {ans}')
    answers['九、句子补全'] = ans_list

    add_answer_section(doc, answers)
    doc.save(output_path)
    print(f'已生成: {output_path}')


# ============ 生成第6天文档函数 ============

def generate_day6_doc(version, output_path):
    """生成第6天语法专项（一般过去时+一般将来时）练习题"""
    doc = Document()

    # 根据版本确定题目数量
    if version == '简洁版':
        past_time = 4
        past_regular, past_irregular = 6, 8
        future_time = 4
        will_sent, going_to = 4, 4
        tense_compare = 3
        tense_choose, fill_tense = 5, 6
        common_errors, complete_tense = 5, 5
    elif version == '完整版':
        past_time = 8
        past_regular, past_irregular = 12, 15
        future_time = 6
        will_sent, going_to = 6, 6
        tense_compare = 3
        tense_choose, fill_tense = 10, 12
        common_errors, complete_tense = 10, 10
    else:  # 充实版
        past_time = 12
        past_regular, past_irregular = 20, 30
        future_time = 10
        will_sent, going_to = 10, 10
        tense_compare = 3
        tense_choose, fill_tense = 15, 20
        common_errors, complete_tense = 15, 15

    add_title(doc, '小学六年级英语练习题【语法专项：一般过去时+一般将来时】', f'（基础+提升）第6天 - {version}')

    answers = {}

    # ========== 第一部分：时态梳理 ==========
    add_section_title(doc, '第一部分：时态梳理（25分钟）')

    # 一、一般过去时标志词
    add_question_header(doc, '一', '一般过去时标志词（写出下列标志词的中文意思）')
    ans_list = []
    for i, (word, meaning, tip) in enumerate(DAY6_PAST_TIME_WORDS[:past_time], 1):
        add_question(doc, f'{i}. {word} → _______')
        ans_list.append(f'{i}. {meaning}（{tip}）')
    answers['一、一般过去时标志词'] = ans_list

    # 二、动词过去式规则变化
    add_question_header(doc, '二', '动词过去式规则变化（写出下列动词的过去式）')
    ans_list = []
    for i, (verb, past, rule) in enumerate(DAY6_PAST_REGULAR[:past_regular], 1):
        add_question(doc, f'{i}. {verb} → _______')
        ans_list.append(f'{i}. {past}（{rule}）')
    answers['二、动词过去式规则变化'] = ans_list

    # 三、动词过去式不规则变化
    add_question_header(doc, '三', '动词过去式不规则变化（写出下列动词的过去式）')
    ans_list = []
    for i, (verb, past, rule) in enumerate(DAY6_PAST_IRREGULAR[:past_irregular], 1):
        add_question(doc, f'{i}. {verb} → _______')
        ans_list.append(f'{i}. {past}（{rule}）')
    answers['三、动词过去式不规则变化'] = ans_list

    # 四、一般将来时标志词
    add_question_header(doc, '四', '一般将来时标志词（写出下列标志词的中文意思）')
    ans_list = []
    for i, (word, meaning, tip) in enumerate(DAY6_FUTURE_TIME_WORDS[:future_time], 1):
        add_question(doc, f'{i}. {word} → _______')
        ans_list.append(f'{i}. {meaning}（{tip}）')
    answers['四、一般将来时标志词'] = ans_list

    # 五、will 句型练习
    add_question_header(doc, '五', 'will 句型练习（翻译下列句子）')
    ans_list = []
    for i, (en, cn, rule) in enumerate(DAY6_WILL_SENTENCES[:will_sent], 1):
        add_question(doc, f'{i}. {en}')
        add_question(doc, '   翻译：_______________________')
        ans_list.append(f'{i}. {cn}（{rule}）')
    answers['五、will 句型练习'] = ans_list

    # 六、be going to 句型练习
    add_question_header(doc, '六', 'be going to 句型练习（翻译下列句子）')
    ans_list = []
    for i, (en, cn, rule) in enumerate(DAY6_BE_GOING_TO[:going_to], 1):
        add_question(doc, f'{i}. {en}')
        add_question(doc, '   翻译：_______________________')
        ans_list.append(f'{i}. {cn}（{rule}）')
    answers['六、be going to 句型练习'] = ans_list

    # 七、三大基础时态对比
    add_question_header(doc, '七', '三大基础时态对比（填写表格）')
    add_question(doc, '请根据例句总结三大时态的用法和标志词：', indent=False)
    doc.add_paragraph()
    add_question(doc, '| 时态 | 例句 | 用法 | 常见标志词 |')
    add_question(doc, '|------|------|------|------------|')
    ans_list = []
    for tense, example, usage, words in DAY6_TENSE_COMPARE[:tense_compare]:
        add_question(doc, f'| {tense} | {example} | _______ | _______ |')
        ans_list.append(f'{tense}：{usage}，标志词：{words}')
    answers['七、三大基础时态对比'] = ans_list

    # ========== 第二部分：例题练习 ==========
    add_section_title(doc, '第二部分：例题练习（15分钟）')

    # 八、时态辨析选择题
    add_question_header(doc, '八', '时态辨析选择题（选择正确答案）')
    ans_list = []
    for i, (sent, options, ans, reason) in enumerate(DAY6_TENSE_CHOOSE[:tense_choose], 1):
        add_question(doc, f'{i}. {sent}')
        add_question(doc, f'   {options}')
        ans_list.append(f'{i}. {ans}（{reason}）')
    answers['八、时态辨析选择题'] = ans_list

    # 九、用所给词适当形式填空
    add_question_header(doc, '九', '用所给词适当形式填空')
    ans_list = []
    for i, (sent, ans, reason) in enumerate(DAY6_FILL_TENSE[:fill_tense], 1):
        add_question(doc, f'{i}. {sent}')
        ans_list.append(f'{i}. {ans}（{reason}）')
    answers['九、用所给词适当形式填空'] = ans_list

    # ========== 第三部分：错题复盘 ==========
    add_section_title(doc, '第三部分：错题复盘（10分钟）')

    # 十、时态易错题整理
    add_question_header(doc, '十', '时态易错题整理（找出错误并改正）')
    ans_list = []
    for i, (sent, correction, reason) in enumerate(DAY6_COMMON_ERRORS[:common_errors], 1):
        add_question(doc, f'{i}. {sent}')
        add_question(doc, '   错误：_______ → 改正：_______')
        ans_list.append(f'{i}. {correction}（{reason}）')
    answers['十、时态易错题整理'] = ans_list

    # 十一、句子补全（三大时态综合）
    add_question_header(doc, '十一', '句子补全（根据中文提示完成英文句子）')
    ans_list = []
    for i, (cn, en, ans) in enumerate(DAY6_COMPLETE_TENSE[:complete_tense], 1):
        add_question(doc, f'{i}. {cn}')
        add_question(doc, f'   {en}')
        ans_list.append(f'{i}. {ans}')
    answers['十一、句子补全'] = ans_list

    add_answer_section(doc, answers)
    doc.save(output_path)
    print(f'已生成: {output_path}')


# ============ 生成第7天文档函数 ============

def generate_day7_doc(version, output_path):
    """生成第7天语法专项（介词+代词+名词单复数）练习题"""
    doc = Document()

    # 根据版本确定题目数量
    if version == '简洁版':
        prep_in, prep_on, prep_at, prep_by = 4, 4, 4, 4
        pronoun_fill, possessive_fill = 5, 5
        noun_regular, noun_irregular = 8, 5
        time_prep, demonstrative = 5, 4
        pronoun_errors = 4
        prep_fill, pronoun_replace, noun_convert = 5, 4, 5
    elif version == '完整版':
        prep_in, prep_on, prep_at, prep_by = 8, 8, 8, 6
        pronoun_fill, possessive_fill = 10, 10
        noun_regular, noun_irregular = 15, 8
        time_prep, demonstrative = 10, 8
        pronoun_errors = 8
        prep_fill, pronoun_replace, noun_convert = 10, 8, 10
    else:  # 充实版
        prep_in, prep_on, prep_at, prep_by = 12, 12, 12, 10
        pronoun_fill, possessive_fill = 15, 15
        noun_regular, noun_irregular = 27, 12
        time_prep, demonstrative = 15, 12
        pronoun_errors = 12
        prep_fill, pronoun_replace, noun_convert = 15, 12, 15

    add_title(doc, '小学六年级英语练习题【语法专项：介词+代词+名词单复数】', f'（基础+提升）第7天 - {version}')

    answers = {}

    # ========== 第一部分：核心语法 ==========
    add_section_title(doc, '第一部分：核心语法（20分钟）')

    # 一、介词 in 的用法
    add_question_header(doc, '一', '介词 in 的用法（写出下列短语的中文意思）')
    ans_list = []
    for i, (phrase, meaning, usage) in enumerate(DAY7_PREPOSITION_IN[:prep_in], 1):
        add_question(doc, f'{i}. {phrase} → _______')
        ans_list.append(f'{i}. {meaning}（{usage}）')
    answers['一、介词 in 的用法'] = ans_list

    # 二、介词 on 的用法
    add_question_header(doc, '二', '介词 on 的用法（写出下列短语的中文意思）')
    ans_list = []
    for i, (phrase, meaning, usage) in enumerate(DAY7_PREPOSITION_ON[:prep_on], 1):
        add_question(doc, f'{i}. {phrase} → _______')
        ans_list.append(f'{i}. {meaning}（{usage}）')
    answers['二、介词 on 的用法'] = ans_list

    # 三、介词 at 的用法
    add_question_header(doc, '三', '介词 at 的用法（写出下列短语的中文意思）')
    ans_list = []
    for i, (phrase, meaning, usage) in enumerate(DAY7_PREPOSITION_AT[:prep_at], 1):
        add_question(doc, f'{i}. {phrase} → _______')
        ans_list.append(f'{i}. {meaning}（{usage}）')
    answers['三、介词 at 的用法'] = ans_list

    # 四、介词 by 的用法
    add_question_header(doc, '四', '介词 by 的用法（写出下列短语的中文意思）')
    ans_list = []
    for i, (phrase, meaning, usage) in enumerate(DAY7_PREPOSITION_BY[:prep_by], 1):
        add_question(doc, f'{i}. {phrase} → _______')
        ans_list.append(f'{i}. {meaning}（{usage}）')
    answers['四、介词 by 的用法'] = ans_list

    # 五、人称代词填空
    add_question_header(doc, '五', '人称代词填空（选择正确的人称代词）')
    add_question(doc, '人称代词表：I-me, you-you, he-him, she-her, it-it, we-us, they-them', indent=False)
    doc.add_paragraph()
    ans_list = []
    for i, (sent, ans, reason) in enumerate(DAY7_PRONOUN_FILL[:pronoun_fill], 1):
        add_question(doc, f'{i}. {sent}')
        ans_list.append(f'{i}. {ans}（{reason}）')
    answers['五、人称代词填空'] = ans_list

    # 六、物主代词填空
    add_question_header(doc, '六', '物主代词填空（选择正确的物主代词）')
    add_question(doc, '物主代词表：my-mine, your-yours, his-his, her-hers, its-its, our-ours, their-theirs', indent=False)
    doc.add_paragraph()
    ans_list = []
    for i, (sent, ans, reason) in enumerate(DAY7_POSSESSIVE_FILL[:possessive_fill], 1):
        add_question(doc, f'{i}. {sent}')
        ans_list.append(f'{i}. {ans}（{reason}）')
    answers['六、物主代词填空'] = ans_list

    # 七、名词单复数规则变化
    add_question_header(doc, '七', '名词单复数规则变化（写出下列名词的复数形式）')
    ans_list = []
    for i, (singular, plural, rule) in enumerate(DAY7_NOUN_PLURAL_REGULAR[:noun_regular], 1):
        add_question(doc, f'{i}. {singular} → _______')
        ans_list.append(f'{i}. {plural}（{rule}）')
    answers['七、名词单复数规则变化'] = ans_list

    # 八、名词单复数不规则变化
    add_question_header(doc, '八', '名词单复数不规则变化（写出下列名词的复数形式）')
    ans_list = []
    for i, (singular, plural, rule) in enumerate(DAY7_NOUN_PLURAL_IRREGULAR[:noun_irregular], 1):
        add_question(doc, f'{i}. {singular} → _______')
        ans_list.append(f'{i}. {plural}（{rule}）')
    answers['八、名词单复数不规则变化'] = ans_list

    # ========== 第二部分：易混辨析 ==========
    add_section_title(doc, '第二部分：易混辨析（15分钟）')

    # 九、in/on/at 时间用法辨析
    add_question_header(doc, '九', 'in/on/at 时间用法辨析（填入正确的介词）')
    ans_list = []
    for i, (sent, ans, reason) in enumerate(DAY7_TIME_PREPOSITION[:time_prep], 1):
        add_question(doc, f'{i}. {sent}')
        ans_list.append(f'{i}. {ans}（{reason}）')
    answers['九、in/on/at 时间用法辨析'] = ans_list

    # 十、this/that/these/those 辨析
    add_question_header(doc, '十', 'this/that/these/those 辨析（填入正确的指示代词）')
    ans_list = []
    for i, (sent, ans, reason) in enumerate(DAY7_DEMONSTRATIVE[:demonstrative], 1):
        add_question(doc, f'{i}. {sent}')
        ans_list.append(f'{i}. {ans}（{reason}）')
    answers['十、this/that/these/those 辨析'] = ans_list

    # 十一、代词易错题
    add_question_header(doc, '十一', '代词易错题（找出错误并改正）')
    ans_list = []
    for i, (sent, correction, reason) in enumerate(DAY7_PRONOUN_ERRORS[:pronoun_errors], 1):
        add_question(doc, f'{i}. {sent}')
        add_question(doc, '   错误：_______ → 改正：_______')
        ans_list.append(f'{i}. {correction}（{reason}）')
    answers['十一、代词易错题'] = ans_list

    # ========== 第三部分：综合练习 ==========
    add_section_title(doc, '第三部分：综合练习（15分钟）')

    # 十二、介词填空综合练习
    add_question_header(doc, '十二', '介词填空综合练习（填入正确的介词 in/on/at/by）')
    ans_list = []
    for i, (sent, ans, reason) in enumerate(DAY7_PREPOSITION_FILL[:prep_fill], 1):
        add_question(doc, f'{i}. {sent}')
        ans_list.append(f'{i}. {ans}（{reason}）')
    answers['十二、介词填空综合练习'] = ans_list

    # 十三、代词替换练习
    add_question_header(doc, '十三', '代词替换练习（用代词替换划线部分）')
    ans_list = []
    for i, (original, replaced, tip) in enumerate(DAY7_PRONOUN_REPLACE[:pronoun_replace], 1):
        add_question(doc, f'{i}. {original}')
        add_question(doc, '   → _______________________')
        ans_list.append(f'{i}. {replaced}（{tip}）')
    answers['十三、代词替换练习'] = ans_list

    # 十四、名词单复数转换练习
    add_question_header(doc, '十四', '名词单复数转换练习（写出正确的复数形式）')
    ans_list = []
    for i, (sent, ans, rule) in enumerate(DAY7_NOUN_CONVERT[:noun_convert], 1):
        add_question(doc, f'{i}. {sent}')
        ans_list.append(f'{i}. {ans}（{rule}）')
    answers['十四、名词单复数转换练习'] = ans_list

    add_answer_section(doc, answers)
    doc.save(output_path)
    print(f'已生成: {output_path}')


# ============ 生成第8天文档函数 ============

def generate_day8_doc(version, output_path):
    """生成第8天阅读专项（完形填空+阅读理解基础）练习题"""
    doc = Document()

    # 根据版本确定内容
    if version == '简洁版':
        tips_count = 3
        cloze_count = 1  # 完形填空篇数
        reading_count = 1  # 阅读理解篇数
        guess_count = 3
    elif version == '完整版':
        tips_count = 5
        cloze_count = 1
        reading_count = 2
        guess_count = 4
    else:  # 充实版
        tips_count = 7
        cloze_count = 2
        reading_count = 2
        guess_count = 6

    add_title(doc, '小学六年级英语练习题【阅读专项：完形填空+阅读理解】', f'（基础+提升）第8天 - {version}')

    answers = {}

    # ========== 第一部分：阅读技巧 ==========
    add_section_title(doc, '第一部分：阅读技巧（20分钟）')

    # 一、阅读技巧知识点
    add_question_header(doc, '一', '阅读技巧知识点（请认真阅读并记忆）')
    for i, (tip, explanation, category) in enumerate(DAY8_READING_TIPS[:tips_count], 1):
        add_question(doc, f'{i}. 【{category}】{tip}')
        add_question(doc, f'   说明：{explanation}')
    answers['一、阅读技巧知识点'] = ['请认真阅读并记忆以上技巧']

    # 二、猜词义练习
    add_question_header(doc, '二', '猜词义练习（根据上下文猜测划线单词的意思）')
    ans_list = []
    for i, (sent, word, meaning, reason) in enumerate(DAY8_GUESS_WORD[:guess_count], 1):
        add_question(doc, f'{i}. {sent}')
        add_question(doc, f'   "{word}" 的意思是：_______')
        ans_list.append(f'{i}. {meaning}（{reason}）')
    answers['二、猜词义练习'] = ans_list

    # ========== 第二部分：题型练习 ==========
    add_section_title(doc, '第二部分：题型练习（25分钟）')

    # 三、完形填空
    add_question_header(doc, '三', '完形填空（阅读短文，选择最佳答案）')
    ans_list = []
    cloze_data = [DAY8_CLOZE_1, DAY8_CLOZE_2][:cloze_count]
    q_num = 1
    for cloze in cloze_data:
        add_question(doc, f'【{cloze["title"]}】', indent=False)
        for line in cloze['passage'].split('\n'):
            if line.strip():
                add_question(doc, line)
        doc.add_paragraph()
        for q, options, ans, reason in cloze['questions']:
            add_question(doc, f'{q_num}. {q}')
            add_question(doc, f'   {options}')
            ans_list.append(f'{q_num}. {ans}（{reason}）')
            q_num += 1
        doc.add_paragraph()
    answers['三、完形填空'] = ans_list

    # 四、阅读理解
    add_question_header(doc, '四', '阅读理解（阅读短文，选择最佳答案）')
    ans_list = []
    reading_data = [DAY8_READING_1, DAY8_READING_2][:reading_count]
    q_num = 1
    for reading in reading_data:
        add_question(doc, f'【{reading["title"]}】', indent=False)
        for line in reading['passage'].split('\n'):
            if line.strip():
                add_question(doc, line)
        doc.add_paragraph()
        for q, options, ans, reason in reading['questions']:
            add_question(doc, f'{q_num}. {q}')
            add_question(doc, f'   {options}')
            ans_list.append(f'{q_num}. {ans}（{reason}）')
            q_num += 1
        doc.add_paragraph()
    answers['四、阅读理解'] = ans_list

    # ========== 第三部分：讲解复盘 ==========
    add_section_title(doc, '第三部分：讲解复盘（5分钟）')

    add_question_header(doc, '五', '错题分析与技巧总结')
    add_question(doc, '完形填空技巧：', indent=False)
    add_question(doc, '1. 先通读全文，了解大意')
    add_question(doc, '2. 注意固定搭配和习惯用法')
    add_question(doc, '3. 根据上下文逻辑选择答案')
    add_question(doc, '4. 做完后再读一遍，检查是否通顺')
    doc.add_paragraph()
    add_question(doc, '阅读理解技巧：', indent=False)
    add_question(doc, '1. 先看题目，带着问题读文章')
    add_question(doc, '2. 圈画关键词：时间、人物、数字、地点')
    add_question(doc, '3. 答案一定在原文中能找到依据')
    add_question(doc, '4. 注意转折词后面的内容')
    answers['五、错题分析与技巧总结'] = ['请认真阅读以上技巧总结']

    add_answer_section(doc, answers)
    doc.save(output_path)
    print(f'已生成: {output_path}')


# ============ 生成第9天文档函数 ============

def generate_day9_doc(version, output_path):
    """生成第9天写作专项（小作文·审题+句型+书写）练习题"""
    doc = Document()

    # 根据版本确定内容
    if version == '简洁版':
        tips_count = 3
        opening_count, body_count, ending_count = 4, 5, 4
        errors_count = 4
        topics_count = 1
    elif version == '完整版':
        tips_count = 4
        opening_count, body_count, ending_count = 6, 7, 6
        errors_count = 6
        topics_count = 2
    else:  # 充实版
        tips_count = 5
        opening_count, body_count, ending_count = 8, 10, 8
        errors_count = 8
        topics_count = 2

    add_title(doc, '小学六年级英语练习题【写作专项：审题+句型+书写】', f'（基础+提升）第9天 - {version}')

    answers = {}

    # ========== 第一部分：审题技巧 ==========
    add_section_title(doc, '第一部分：审题技巧（15分钟）')

    # 一、审题技巧
    add_question_header(doc, '一', '审题技巧（请认真阅读并记忆）')
    for i, (tip, explanation, category) in enumerate(DAY9_WRITING_TIPS[:tips_count], 1):
        add_question(doc, f'{i}. 【{tip}】{explanation}')
    answers['一、审题技巧'] = ['请认真阅读并记忆以上技巧']

    # 二、写作框架
    add_question_header(doc, '二', '写作框架（开头-中间-结尾）')
    add_question(doc, '一篇好的作文应该包含三个部分：', indent=False)
    add_question(doc, '1. 开头（Opening）：引出话题，点明主题')
    add_question(doc, '2. 中间（Body）：详细描述，展开内容')
    add_question(doc, '3. 结尾（Ending）：总结感受，呼应开头')
    answers['二、写作框架'] = ['开头-中间-结尾三段式结构']

    # ========== 第二部分：句型积累 ==========
    add_section_title(doc, '第二部分：句型积累（10分钟）')

    # 三、常用开头句
    add_question_header(doc, '三', '常用开头句（请背诵并仿写）')
    ans_list = []
    for i, (en, cn, usage) in enumerate(DAY9_OPENING_SENTENCES[:opening_count], 1):
        add_question(doc, f'{i}. {en}')
        add_question(doc, f'   中文：{cn}（{usage}）')
        ans_list.append(f'{i}. {en} - {cn}')
    answers['三、常用开头句'] = ans_list

    # 四、常用中间句
    add_question_header(doc, '四', '常用中间句（描述句型）')
    ans_list = []
    for i, (en, cn, usage) in enumerate(DAY9_BODY_SENTENCES[:body_count], 1):
        add_question(doc, f'{i}. {en}')
        add_question(doc, f'   中文：{cn}（{usage}）')
        ans_list.append(f'{i}. {en} - {cn}')
    answers['四、常用中间句'] = ans_list

    # 五、常用结尾句
    add_question_header(doc, '五', '常用结尾句')
    ans_list = []
    for i, (en, cn, usage) in enumerate(DAY9_ENDING_SENTENCES[:ending_count], 1):
        add_question(doc, f'{i}. {en}')
        add_question(doc, f'   中文：{cn}（{usage}）')
        ans_list.append(f'{i}. {en} - {cn}')
    answers['五、常用结尾句'] = ans_list

    # 六、写作常见错误
    add_question_header(doc, '六', '写作常见错误（找出错误并改正）')
    ans_list = []
    for i, (wrong, correct, reason) in enumerate(DAY9_WRITING_ERRORS[:errors_count], 1):
        add_question(doc, f'{i}. 错误：{wrong}')
        add_question(doc, '   改正：_______________________')
        ans_list.append(f'{i}. {correct}（{reason}）')
    answers['六、写作常见错误'] = ans_list

    # ========== 第三部分：写作实战 ==========
    add_section_title(doc, '第三部分：写作实战（25分钟）')

    topics = [DAY9_WRITING_TOPIC_1, DAY9_WRITING_TOPIC_2][:topics_count]
    ans_list = []
    for idx, topic in enumerate(topics, 7):
        add_question_header(doc, '七' if idx == 7 else '八', f'写作题目：{topic["title"]}')
        add_question(doc, f'要求：{topic["requirements"]}', indent=False)
        add_question(doc, '提示：', indent=False)
        for hint in topic['hints']:
            add_question(doc, f'• {hint}')
        add_question(doc, '写作框架：', indent=False)
        add_question(doc, f'• 开头：{topic["outline"]["opening"]}')
        add_question(doc, f'• 中间：{topic["outline"]["body"]}')
        add_question(doc, f'• 结尾：{topic["outline"]["ending"]}')
        doc.add_paragraph()
        add_question(doc, '请在下面写作：', indent=False)
        add_question(doc, '_' * 50)
        add_question(doc, '_' * 50)
        add_question(doc, '_' * 50)
        add_question(doc, '_' * 50)
        add_question(doc, '_' * 50)
        doc.add_paragraph()
        ans_list.append(f'【{topic["title"]} 范文】')
        ans_list.append(topic['sample'])

    answers['七、写作实战'] = ans_list

    add_answer_section(doc, answers)
    doc.save(output_path)
    print(f'已生成: {output_path}')


# ============ 生成第10天文档函数 ============

def generate_day10_doc(version, output_path):
    """生成第10天综合模考+错题复盘+知识梳理练习题"""
    doc = Document()

    # 根据版本确定内容
    if version == '简洁版':
        vocab_count = 8
        grammar_count = 10
    elif version == '完整版':
        vocab_count = 12
        grammar_count = 15
    else:  # 充实版
        vocab_count = 15
        grammar_count = 20

    add_title(doc, '小学六年级英语练习题【综合模考+错题复盘+知识梳理】', f'（基础+提升）第10天 - {version}')

    answers = {}

    # ========== 第一部分：综合模考 ==========
    add_section_title(doc, '第一部分：综合模考（30分钟）')

    # 一、词汇部分
    add_question_header(doc, '一', '词汇部分（填写正确的单词或短语）')
    ans_list = []
    for i, (question, ans, q_type, tip) in enumerate(DAY10_VOCAB_TEST[:vocab_count], 1):
        add_question(doc, f'{i}. {question}')
        ans_list.append(f'{i}. {ans}（{tip}）')
    answers['一、词汇部分'] = ans_list

    # 二、语法部分
    add_question_header(doc, '二', '语法部分（用所给词的正确形式填空或选择）')
    ans_list = []
    for i, (question, ans, tip) in enumerate(DAY10_GRAMMAR_TEST[:grammar_count], 1):
        add_question(doc, f'{i}. {question}')
        ans_list.append(f'{i}. {ans}（{tip}）')
    answers['二、语法部分'] = ans_list

    # 三、阅读理解
    add_question_header(doc, '三', '阅读理解（阅读短文，选择最佳答案）')
    reading = DAY10_READING_TEST
    add_question(doc, f'【{reading["title"]}】', indent=False)
    for line in reading['passage'].split('\n'):
        if line.strip():
            add_question(doc, line)
    doc.add_paragraph()
    ans_list = []
    for i, (q, options, ans, reason) in enumerate(reading['questions'], 1):
        add_question(doc, f'{i}. {q}')
        add_question(doc, f'   {options}')
        ans_list.append(f'{i}. {ans}（{reason}）')
    answers['三、阅读理解'] = ans_list

    # 四、写作
    add_question_header(doc, '四', f'写作：{DAY10_WRITING_TEST["title"]}')
    add_question(doc, f'要求：{DAY10_WRITING_TEST["requirements"]}', indent=False)
    add_question(doc, '提示：', indent=False)
    for hint in DAY10_WRITING_TEST['hints']:
        add_question(doc, f'• {hint}')
    doc.add_paragraph()
    add_question(doc, '请在下面写作：', indent=False)
    add_question(doc, '_' * 50)
    add_question(doc, '_' * 50)
    add_question(doc, '_' * 50)
    add_question(doc, '_' * 50)
    answers['四、写作'] = ['请参考第9天范文格式自行评分']

    # ========== 第二部分：快速批改 ==========
    add_section_title(doc, '第二部分：快速批改（10分钟）')

    add_question_header(doc, '五', '自批试卷（对照答案批改，标注错题类型）')
    add_question(doc, '错题类型分类：', indent=False)
    add_question(doc, '• 词汇类错误：短语搭配、单词拼写、名词复数等')
    add_question(doc, '• 语法类错误：时态、人称代词、介词等')
    add_question(doc, '• 阅读类错误：理解偏差、定位不准等')
    doc.add_paragraph()
    add_question(doc, '我的错题统计：', indent=False)
    add_question(doc, '词汇类错误：_______ 题')
    add_question(doc, '语法类错误：_______ 题')
    add_question(doc, '阅读类错误：_______ 题')
    answers['五、自批试卷'] = ['请自行统计错题数量']

    # ========== 第三部分：复盘梳理 ==========
    add_section_title(doc, '第三部分：复盘梳理（10分钟）')

    # 六、10天学习要点回顾
    add_question_header(doc, '六', '10天学习要点回顾')
    for day, topic, content in DAY10_REVIEW_POINTS:
        add_question(doc, f'【{day}】{topic}')
        add_question(doc, f'   {content}')
    answers['六、10天学习要点回顾'] = ['请认真复习以上知识点']

    # 七、易错点汇总
    add_question_header(doc, '七', '易错点汇总（重点巩固）')
    ans_list = []
    for i, (error_type, example, tip) in enumerate(DAY10_ERROR_SUMMARY, 1):
        add_question(doc, f'{i}. 【{error_type}】')
        add_question(doc, f'   例句：{example}')
        add_question(doc, f'   注意：{tip}')
        ans_list.append(f'{i}. {error_type}：{tip}')
    answers['七、易错点汇总'] = ans_list

    # 八、后续学习建议
    add_question_header(doc, '八', '后续学习建议')
    add_question(doc, '1. 每天复习错题本，巩固薄弱知识点')
    add_question(doc, '2. 坚持每天背诵5-10个单词和短语')
    add_question(doc, '3. 每周完成1-2篇阅读理解练习')
    add_question(doc, '4. 每周写1篇小作文，积累常用句型')
    add_question(doc, '5. 多听多读，培养语感')
    answers['八、后续学习建议'] = ['请认真执行以上学习计划']

    add_answer_section(doc, answers)
    doc.save(output_path)
    print(f'已生成: {output_path}')


# ============ 主函数 ============

def main():
    """生成所有练习题文档"""
    import os

    # 获取脚本所在目录
    script_dir = os.path.dirname(os.path.abspath(__file__))

    versions = ['简洁版', '完整版', '充实版']

    print('开始生成英语练习题文档...\n')

    # 生成第2天文档
    for version in versions:
        filename = f'第2天_词汇基础_{version}.docx'
        output_path = os.path.join(script_dir, filename)
        generate_day2_doc(version, output_path)

    # 生成第3天文档
    for version in versions:
        filename = f'第3天_句型基础_{version}.docx'
        output_path = os.path.join(script_dir, filename)
        generate_day3_doc(version, output_path)

    # 生成第4天文档
    for version in versions:
        filename = f'第4天_特殊疑问句与情景交际_{version}.docx'
        output_path = os.path.join(script_dir, filename)
        generate_day4_doc(version, output_path)

    # 生成第5天文档
    for version in versions:
        filename = f'第5天_语法专项_一般现在时_{version}.docx'
        output_path = os.path.join(script_dir, filename)
        generate_day5_doc(version, output_path)

    # 生成第6天文档
    for version in versions:
        filename = f'第6天_语法专项_一般过去时与将来时_{version}.docx'
        output_path = os.path.join(script_dir, filename)
        generate_day6_doc(version, output_path)

    # 生成第7天文档
    for version in versions:
        filename = f'第7天_语法专项_介词代词名词单复数_{version}.docx'
        output_path = os.path.join(script_dir, filename)
        generate_day7_doc(version, output_path)

    # 生成第8天文档
    for version in versions:
        filename = f'第8天_阅读专项_完形填空与阅读理解_{version}.docx'
        output_path = os.path.join(script_dir, filename)
        generate_day8_doc(version, output_path)

    # 生成第9天文档
    for version in versions:
        filename = f'第9天_写作专项_审题句型书写_{version}.docx'
        output_path = os.path.join(script_dir, filename)
        generate_day9_doc(version, output_path)

    # 生成第10天文档
    for version in versions:
        filename = f'第10天_综合模考与知识梳理_{version}.docx'
        output_path = os.path.join(script_dir, filename)
        generate_day10_doc(version, output_path)

    print('\n所有文档生成完成！')


if __name__ == '__main__':
    main()
